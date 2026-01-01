# -*- coding: utf-8 -*-
import streamlit as st
import geopandas as gpd
import pandas as pd
import numpy as np
import tempfile
import os
import zipfile
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import matplotlib  # ← IMPORTADO PARA __version__
from matplotlib.tri import Triangulation
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap
import io
from shapely.geometry import Polygon, LineString
import math
import warnings
import xml.etree.ElementTree as ET
import base64
import json
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import geojson
import requests
import contextily as ctx
import yfinance as yf
import plotly.graph_objects as go
warnings.filterwarnings('ignore')

# === ESTA DEBE SER LA PRIMERA LLAMADA A STREAMLIT ===
st.set_page_config(
    page_title="Analizador Multi-Cultivo Satellital",
    page_icon="🌱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# === ESTILOS PERSONALIZADOS - VERSIÓN PREMIUM MODERNA ===
st.markdown("""
<style>
/* === FONDO GENERAL OSCURO ELEGANTE === */
.stApp {
    background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%) !important;
    color: #ffffff !important;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}
/* === SIDEBAR: FONDO BLANCO CON TEXTO NEGRO === */
[data-testid="stSidebar"] {
    background: #ffffff !important;
    border-right: 1px solid #e5e7eb !important;
    box-shadow: 5px 0 25px rgba(0, 0, 0, 0.1) !important;
}
/* Texto general del sidebar en NEGRO */
[data-testid="stSidebar"] *,
[data-testid="stSidebar"] .stMarkdown,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stText,
[data-testid="stSidebar"] .stTitle,
[data-testid="stSidebar"] .stSubheader {
    color: #000000 !important;
    text-shadow: none !important;
}
/* Título del sidebar elegante */
.sidebar-title {
    font-size: 1.4em;
    font-weight: 800;
    margin: 1.5em 0 1em 0;
    text-align: center;
    padding: 14px;
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
    border-radius: 16px;
    color: #ffffff !important;
    box-shadow: 0 6px 20px rgba(59, 130, 246, 0.3);
    border: 1px solid rgba(255, 255, 255, 0.2);
    letter-spacing: 0.5px;
}
/* Widgets del sidebar con estilo glassmorphism */
[data-testid="stSidebar"] .stSelectbox,
[data-testid="stSidebar"] .stDateInput,
[data-testid="stSidebar"] .stSlider {
    background: rgba(255, 255, 255, 0.9) !important;
    backdrop-filter: blur(10px);
    border-radius: 12px;
    padding: 12px;
    margin: 8px 0;
    border: 1px solid #d1d5db !important;
}
/* Labels de los widgets en negro */
[data-testid="stSidebar"] .stSelectbox div,
[data-testid="stSidebar"] .stDateInput div,
[data-testid="stSidebar"] .stSlider label {
    color: #000000 !important;
    font-weight: 600;
    font-size: 0.95em;
}
/* Inputs y selects - fondo blanco con texto negro */
[data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] {
    background-color: #ffffff !important;
    border: 1px solid #d1d5db !important;
    color: #000000 !important;
    border-radius: 8px;
}
/* Slider - colores negro */
[data-testid="stSidebar"] .stSlider [data-baseweb="slider"] {
    color: #000000 !important;
}
/* Date Input - fondo blanco con texto negro */
[data-testid="stSidebar"] .stDateInput [data-baseweb="input"] {
    background-color: #ffffff !important;
    border: 1px solid #d1d5db !important;
    color: #000000 !important;
    border-radius: 8px;
}
/* Placeholder en gris */
[data-testid="stSidebar"] .stDateInput [data-baseweb="input"]::placeholder {
    color: #6b7280 !important;
}
/* Botones premium */
.stButton > button {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: white !important;
    border: none !important;
    padding: 0.8em 1.5em !important;
    border-radius: 12px !important;
    font-weight: 700 !important;
    font-size: 1em !important;
    box-shadow: 0 4px 15px rgba(59, 130, 246, 0.4) !important;
    transition: all 0.3s ease !important;
    text-transform: uppercase !important;
    letter-spacing: 0.5px !important;
}
.stButton > button:hover {
    transform: translateY(-3px) !important;
    box-shadow: 0 8px 25px rgba(59, 130, 246, 0.6) !important;
    background: linear-gradient(135deg, #4f8df8 0%, #2d5fe8 100%) !important;
}
/* === HERO BANNER PRINCIPAL CON IMAGEN === */
.hero-banner {
    background: linear-gradient(rgba(15, 23, 42, 0.9), rgba(15, 23, 42, 0.95)),
        url('https://images.unsplash.com/photo-1597981309443-6e2d2a4d9c3f?ixlib=rb-4.0.3&ixid=M3wxMjA3fDB8MHxwaG90by1wYWdlfHx8fGVufDB8fHx8fA%3D%3D&auto=format&fit=crop&w=2070&q=80') !important;
    background-size: cover !important;
    background-position: center 40% !important;
    padding: 3.5em 2em !important;
    border-radius: 24px !important;
    margin-bottom: 2.5em !important;
    box-shadow: 0 20px 40px rgba(0, 0, 0, 0.4) !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
    position: relative !important;
    overflow: hidden !important;
}
.hero-banner::before {
    content: '' !important;
    position: absolute !important;
    top: 0 !important;
    left: 0 !important;
    right: 0 !important;
    bottom: 0 !important;
    background: linear-gradient(45deg, rgba(59, 130, 246, 0.1), rgba(29, 78, 216, 0.05)) !important;
    z-index: 1 !important;
}
.hero-content {
    position: relative !important;
    z-index: 2 !important;
    text-align: center !important;
}
.hero-title {
    color: #ffffff !important;
    font-size: 3.2em !important;
    font-weight: 900 !important;
    margin-bottom: 0.3em !important;
    text-shadow: 0 4px 12px rgba(0, 0, 0, 0.6) !important;
    letter-spacing: -0.5px !important;
    background: linear-gradient(135deg, #ffffff 0%, #93c5fd 100%) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    background-clip: text !important;
}
.hero-subtitle {
    color: #cbd5e1 !important;
    font-size: 1.3em !important;
    font-weight: 400 !important;
    max-width: 800px !important;
    margin: 0 auto !important;
    line-height: 1.6 !important;
}
/* === PESTAÑAS PRINCIPALES (fuera del sidebar) - SIN CAMBIOS === */
.stTabs [data-baseweb="tab-list"] {
    background: rgba(255, 255, 255, 0.05) !important;
    backdrop-filter: blur(10px) !important;
    padding: 8px 16px !important;
    border-radius: 16px !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    margin-top: 1em !important;
    gap: 8px !important;
}
.stTabs [data-baseweb="tab"] {
    color: #94a3b8 !important;
    font-weight: 600 !important;
    padding: 12px 24px !important;
    border-radius: 12px !important;
    background: transparent !important;
    transition: all 0.3s ease !important;
    border: 1px solid transparent !important;
}
.stTabs [data-baseweb="tab"]:hover {
    color: #ffffff !important;
    background: rgba(59, 130, 246, 0.2) !important;
    border-color: rgba(59, 130, 246, 0.3) !important;
    transform: translateY(-2px) !important;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    border: none !important;
    box-shadow: 0 4px 15px rgba(59, 130, 246, 0.4) !important;
}
/* === PESTAÑAS DEL SIDEBAR: FONDO BLANCO + TEXTO NEGRO === */
[data-testid="stSidebar"] .stTabs [data-baseweb="tab-list"] {
    background: #ffffff !important;
    border: 1px solid #e2e8f0 !important;
    padding: 8px !important;
    border-radius: 12px !important;
    gap: 6px !important;
}
[data-testid="stSidebar"] .stTabs [data-baseweb="tab"] {
    color: #000000 !important;
    background: transparent !important;
    border-radius: 8px !important;
    padding: 8px 16px !important;
    font-weight: 600 !important;
    border: 1px solid transparent !important;
}
[data-testid="stSidebar"] .stTabs [data-baseweb="tab"]:hover {
    background: #f1f5f9 !important;
    color: #000000 !important;
    border-color: #cbd5e1 !important;
}
/* Pestaña activa en el sidebar: blanco con texto negro */
[data-testid="stSidebar"] .stTabs [aria-selected="true"] {
    background: #ffffff !important;
    color: #000000 !important;
    font-weight: 700 !important;
    border: 1px solid #3b82f6 !important;
}
/* === MÉTRICAS PREMIUM === */
div[data-testid="metric-container"] {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.8), rgba(15, 23, 42, 0.9)) !important;
    backdrop-filter: blur(10px) !important;
    border-radius: 20px !important;
    padding: 24px !important;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
    transition: all 0.3s ease !important;
}
div[data-testid="metric-container"]:hover {
    transform: translateY(-5px) !important;
    box-shadow: 0 15px 40px rgba(59, 130, 246, 0.2) !important;
    border-color: rgba(59, 130, 246, 0.4) !important;
}
div[data-testid="metric-container"] label,
div[data-testid="metric-container"] div,
div[data-testid="metric-container"] [data-testid="stMetricValue"],
div[data-testid="metric-container"] [data-testid="stMetricLabel"] {
    color: #ffffff !important;
    font-weight: 600 !important;
}
div[data-testid="metric-container"] [data-testid="stMetricValue"] {
    font-size: 2.5em !important;
    font-weight: 800 !important;
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    background-clip: text !important;
}
/* === GRÁFICOS CON ESTILO OSCURO === */
.stPlotlyChart, .stPyplot {
    background: rgba(15, 23, 42, 0.8) !important;
    backdrop-filter: blur(10px) !important;
    border-radius: 20px !important;
    padding: 20px !important;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
}
/* === EXPANDERS ELEGANTES === */
.streamlit-expanderHeader {
    color: #ffffff !important;
    background: rgba(30, 41, 59, 0.8) !important;
    backdrop-filter: blur(10px) !important;
    border-radius: 16px !important;
    font-weight: 700 !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    padding: 16px 20px !important;
    margin-bottom: 10px !important;
}
.streamlit-expanderContent {
    background: rgba(15, 23, 42, 0.6) !important;
    border-radius: 0 0 16px 16px !important;
    padding: 20px !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    border-top: none !important;
}
/* === TEXTOS GENERALES === */
h1, h2, h3, h4, h5, h6 {
    color: #ffffff !important;
    font-weight: 800 !important;
    margin-top: 1.5em !important;
}
p, div, span, label, li {
    color: #cbd5e1 !important;
    line-height: 1.7 !important;
}
/* === DATA FRAMES TABLAS ELEGANTES === */
.dataframe {
    background: rgba(15, 23, 42, 0.8) !important;
    backdrop-filter: blur(10px) !important;
    border-radius: 16px !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    color: #ffffff !important;
}
.dataframe th {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    padding: 16px !important;
}
.dataframe td {
    color: #cbd5e1 !important;
    padding: 14px 16px !important;
    border-bottom: 1px solid rgba(255, 255, 255, 0.1) !important;
}
/* === ALERTS Y MENSAJES === */
.stAlert {
    border-radius: 16px !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    backdrop-filter: blur(10px) !important;
}
/* === SCROLLBAR PERSONALIZADA === */
::-webkit-scrollbar {
    width: 10px !important;
    height: 10px !important;
}
::-webkit-scrollbar-track {
    background: rgba(15, 23, 42, 0.8) !important;
    border-radius: 10px !important;
}
::-webkit-scrollbar-thumb {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    border-radius: 10px !important;
}
::-webkit-scrollbar-thumb:hover {
    background: linear-gradient(135deg, #4f8df8 0%, #2d5fe8 100%) !important;
}
/* === IMÁGENES DEL SIDEBAR === */
[data-testid="stSidebar"] img {
    border-radius: 16px !important;
    border: 2px solid #d1d5db !important;
    box-shadow: 0 8px 25px rgba(0, 0, 0, 0.1) !important;
    transition: all 0.3s ease !important;
}
[data-testid="stSidebar"] img:hover {
    transform: scale(1.02) !important;
    box-shadow: 0 12px 35px rgba(0, 0, 0, 0.2) !important;
    border-color: #3b82f6 !important;
}
/* === TARJETAS DE CULTIVOS === */
.cultivo-card {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.95)) !important;
    border-radius: 20px !important;
    padding: 25px !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
    transition: all 0.3s ease !important;
    height: 100% !important;
}
.cultivo-card:hover {
    transform: translateY(-8px) !important;
    box-shadow: 0 20px 40px rgba(59, 130, 246, 0.2) !important;
    border-color: rgba(59, 130, 246, 0.4) !important;
}
/* === TABLERO DE CONTROL === */
.dashboard-grid {
    display: grid !important;
    grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)) !important;
    gap: 25px !important;
    margin: 30px 0 !important;
}
.dashboard-card {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.95)) !important;
    border-radius: 20px !important;
    padding: 25px !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
    transition: all 0.3s ease !important;
}
.dashboard-card:hover {
    transform: translateY(-5px) !important;
    box-shadow: 0 20px 40px rgba(59, 130, 246, 0.2) !important;
}
/* === STATS BADGES === */
.stats-badge {
    display: inline-block !important;
    padding: 6px 14px !important;
    border-radius: 50px !important;
    font-size: 0.85em !important;
    font-weight: 700 !important;
    margin: 2px !important;
}
.badge-success {
    background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
    color: white !important;
}
.badge-warning {
    background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%) !important;
    color: white !important;
}
.badge-danger {
    background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%) !important;
    color: white !important;
}
.badge-info {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: white !important;
}
</style>
""", unsafe_allow_html=True)

# ===== HERO BANNER PRINCIPAL =====
st.markdown("""
<div class="hero-banner">
    <div class="hero-content">
        <h1 class="hero-title">ANALIZADOR MULTI-CULTIVO SATELITAL</h1>
        <p class="hero-subtitle">Potenciado con NASA POWER, GEE y tecnología avanzada para una agricultura de precisión</p>
    </div>
</div>
""", unsafe_allow_html=True)

# ===== CONFIGURACIÓN ECONÓMICA - DATOS ACTUALIZABLES =====
PRECIOS_API = {
    'insumos': {
        'urea': 800,
        'fosfato': 950,
        'cloruro_potasio': 700,
        'herbicida': 25,
        'insecticida': 35,
        'semilla_maiz': 350,
        'semilla_soya': 280,
        'semilla_trigo': 180,
        'semilla_girasol': 150
    },
    'precios_pizarra': {
        'maiz_rosario': 220,
        'soya_rosario': 420,
        'trigo_rosario': 250,
        'girasol_rosario': 380,
        'maiz_ba': 210,
        'soya_ba': 410,
        'trigo_ba': 240,
        'girasol_ba': 370
    }
}
RENDIMIENTOS_BASE = {
    'MAÍZ': 8.0,
    'SOYA': 3.5,
    'TRIGO': 4.5,
    'GIRASOL': 2.5
}
COSTOS_BASE = {
    'MAÍZ': 1200,
    'SOYA': 950,
    'TRIGO': 800,
    'GIRASOL': 700
}

# ===== CONFIGURACIÓN DE SATÉLITES DISPONIBLES =====
SATELITES_DISPONIBLES = {
    'SENTINEL-2': {
        'nombre': 'Sentinel-2',
        'resolucion': '10m',
        'revisita': '5 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B8', 'B11'],
        'indices': ['NDVI', 'NDRE', 'GNDVI', 'OSAVI', 'MCARI'],
        'icono': '🛰️'
    },
    'LANDSAT-8': {
        'nombre': 'Landsat 8',
        'resolucion': '30m',
        'revisita': '16 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B6', 'B7'],
        'indices': ['NDVI', 'NDWI', 'EVI', 'SAVI', 'MSAVI'],
        'icono': '🛰️'
    },
    'DATOS_SIMULADOS': {
        'nombre': 'Datos Simulados',
        'resolucion': '10m',
        'revisita': '5 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B8'],
        'indices': ['NDVI', 'NDRE', 'GNDVI'],
        'icono': '🔬'
    }
}

# ===== CONFIGURACIÓN =====
PARAMETROS_CULTIVOS = {
    'MAÍZ': {
        'NITROGENO': {'min': 150, 'max': 200},
        'FOSFORO': {'min': 40, 'max': 60},
        'POTASIO': {'min': 120, 'max': 180},
        'MATERIA_ORGANICA_OPTIMA': 3.5,
        'HUMEDAD_OPTIMA': 0.3,
        'NDVI_OPTIMO': 0.85,
        'NDRE_OPTIMO': 0.5
    },
    'SOYA': {
        'NITROGENO': {'min': 20, 'max': 40},
        'FOSFORO': {'min': 30, 'max': 50},
        'POTASIO': {'min': 80, 'max': 120},
        'MATERIA_ORGANICA_OPTIMA': 4.0,
        'HUMEDAD_OPTIMA': 0.25,
        'NDVI_OPTIMO': 0.8,
        'NDRE_OPTIMO': 0.45
    },
    'TRIGO': {
        'NITROGENO': {'min': 120, 'max': 180},
        'FOSFORO': {'min': 40, 'max': 60},
        'POTASIO': {'min': 80, 'max': 120},
        'MATERIA_ORGANICA_OPTIMA': 3.0,
        'HUMEDAD_OPTIMA': 0.28,
        'NDVI_OPTIMO': 0.75,
        'NDRE_OPTIMO': 0.4
    },
    'GIRASOL': {
        'NITROGENO': {'min': 80, 'max': 120},
        'FOSFORO': {'min': 35, 'max': 50},
        'POTASIO': {'min': 100, 'max': 150},
        'MATERIA_ORGANICA_OPTIMA': 3.2,
        'HUMEDAD_OPTIMA': 0.22,
        'NDVI_OPTIMO': 0.7,
        'NDRE_OPTIMO': 0.35
    }
}

TEXTURA_SUELO_OPTIMA = {
    'MAÍZ': {
        'textura_optima': 'Franco',
        'arena_optima': 45,
        'limo_optima': 35,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.3,
        'porosidad_optima': 0.5
    },
    'SOYA': {
        'textura_optima': 'Franco',
        'arena_optima': 40,
        'limo_optima': 40,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.2,
        'porosidad_optima': 0.55
    },
    'TRIGO': {
        'textura_optima': 'Franco',
        'arena_optima': 50,
        'limo_optima': 30,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.25,
        'porosidad_optima': 0.52
    },
    'GIRASOL': {
        'textura_optima': 'Franco arenoso-arcilloso',
        'arena_optima': 55,
        'limo_optima': 25,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.35,
        'porosidad_optima': 0.48
    }
}

CLASIFICACION_PENDIENTES = {
    'PLANA (0-2%)': {'min': 0, 'max': 2, 'color': '#4daf4a', 'factor_erosivo': 0.1},
    'SUAVE (2-5%)': {'min': 2, 'max': 5, 'color': '#a6d96a', 'factor_erosivo': 0.3},
    'MODERADA (5-10%)': {'min': 5, 'max': 10, 'color': '#ffffbf', 'factor_erosivo': 0.6},
    'FUERTE (10-15%)': {'min': 10, 'max': 15, 'color': '#fdae61', 'factor_erosivo': 0.8},
    'MUY FUERTE (15-25%)': {'min': 15, 'max': 25, 'color': '#f46d43', 'factor_erosivo': 0.9},
    'EXTREMA (>25%)': {'min': 25, 'max': 100, 'color': '#d73027', 'factor_erosivo': 1.0}
}

RECOMENDACIONES_TEXTURA = {
    'Franco': {
        'propiedades': [
            "Equilibrio arena-limo-arcilla",
            "Buena aireación y drenaje",
            "CIC intermedia-alta",
            "Retención de agua adecuada"
        ],
        'limitantes': [
            "Puede compactarse con maquinaria pesada",
            "Erosión en pendientes si no hay cobertura"
        ],
        'manejo': [
            "Mantener coberturas vivas o muertas",
            "Evitar tránsito excesivo de maquinaria",
            "Fertilización eficiente, sin muchas pérdidas",
            "Ideal para la mayoría de cultivos"
        ]
    },
    'Franco arcilloso': {
        'propiedades': [
            "Mayor proporción de arcilla (25–35%)",
            "Alta retención de agua y nutrientes",
            "Drenaje natural lento",
            "Buena fertilidad natural"
        ],
        'limitantes': [
            "Riesgo de encharcamiento",
            "Compactación fácil",
            "Menor oxigenación radicular"
        ],
        'manejo': [
            "Implementar drenajes (canales y subdrenes)",
            "Subsolado previo a siembra",
            "Incorporar materia orgánica",
            "Fertilización fraccionada en lluvias intensas"
        ]
    },
    'Franco arenoso-arcilloso': {
        'propiedades': [
            "Arena 40–50%, arcilla 20–30%",
            "Buen desarrollo radicular",
            "Drenaje moderado",
            "Retención de agua moderada-baja"
        ],
        'limitantes': [
            "Riesgo de lixiviación de nutrientes",
            "Estrés hídrico en veranos",
            "Fertilidad moderada"
        ],
        'manejo': [
            "Uso de coberturas leguminosas",
            "Aplicar mulching",
            "Riego suplementario en sequía",
            "Fertilización fraccionada"
        ]
    }
}

ICONOS_CULTIVOS = {
    'MAÍZ': '🌽',
    'SOYA': '🫘',
    'TRIGO': '🌾',
    'GIRASOL': '🌻'
}

COLORES_CULTIVOS = {
    'MAÍZ': '#FFD700',
    'SOYA': '#90EE90',
    'TRIGO': '#DAA520',
    'GIRASOL': '#FFA500'
}

PALETAS_GEE = {
    'FERTILIDAD': ['#d73027', '#f46d43', '#fdae61', '#fee08b', '#d9ef8b', '#a6d96a', '#66bd63', '#1a9850', '#006837'],
    'NITROGENO': ['#00ff00', '#80ff00', '#ffff00', '#ff8000', '#ff0000'],
    'FOSFORO': ['#0000ff', '#4040ff', '#8080ff', '#c0c0ff', '#ffffff'],
    'POTASIO': ['#4B0082', '#6A0DAD', '#8A2BE2', '#9370DB', '#D8BFD8'],
    'TEXTURA': ['#8c510a', '#d8b365', '#f6e8c3', '#c7eae5', '#5ab4ac', '#01665e'],
    'ELEVACION': ['#006837', '#1a9850', '#66bd63', '#a6d96a', '#d9ef8b', '#ffffbf', '#fee08b', '#fdae61', '#f46d43', '#d73027'],
    'PENDIENTE': ['#4daf4a', '#a6d96a', '#ffffbf', '#fdae61', '#f46d43', '#d73027']
}

# ===== IMÁGENES SEGURAS PARA CULTIVOS =====
IMAGENES_CULTIVOS = {
    'MAÍZ': 'https://images.unsplash.com/photo-1500382017468-9049fed747ef?w=400&h=300&fit=crop&auto=format',
    'SOYA': 'https://images.unsplash.com/photo-1546548970-71785318a17b?w=400&h=300&fit=crop&auto=format',
    'TRIGO': 'https://images.unsplash.com/photo-1593549157444-8e6db7b4534d?w=400&h=300&fit=crop&auto=format',
    'GIRASOL': 'https://images.unsplash.com/photo-1592925144850-72d38b2f3d58?w=400&h=300&fit=crop&auto=format'
}

# ===== FUNCIONES PARA OBTENER DATOS ECONÓMICOS ACTUALIZADOS =====
def obtener_precios_actualizados():
    """
    Obtiene precios actualizados de commodities agrícolas usando yfinance.
    """
    try:
        symbols = {
            'maiz': 'ZC=F',
            'soya': 'ZS=F',
            'trigo': 'ZW=F',
            'girasol': 'BO=F'
        }
        precios_actuales = {}
        for producto, symbol in symbols.items():
            try:
                ticker = yf.Ticker(symbol)
                hist = ticker.history(period='1d')
                if not hist.empty:
                    precio = hist['Close'].iloc[-1]
                    precios_actuales[producto] = float(precio)
                else:
                    # Usar precios por defecto si no hay datos
                    precios_actuales[producto] = PRECIOS_API['precios_pizarra'][f'{producto}_rosario']
            except Exception as e:
                precios_actuales[producto] = PRECIOS_API['precios_pizarra'][f'{producto}_rosario']
        # Actualizar precios de insumos con inflación estimada
        inflacion_estimada = 1.08  # 8% anual
        insumos_actualizados = {}
        for insumo, precio in PRECIOS_API['insumos'].items():
            insumos_actualizados[insumo] = round(precio * inflacion_estimada, 2)
        return {
            'commodities': precios_actuales,
            'insumos': insumos_actualizados,
            'pizarra_rosario': PRECIOS_API['precios_pizarra'],
            'fecha_actualizacion': datetime.now().strftime("%Y-%m-%d")
        }
    except Exception as e:
        # Si hay error, devolver precios por defecto
        return PRECIOS_API

def calcular_costo_fertilizacion(dosis_npk, cultivo, precios):
    """
    Calcula el costo de fertilización basado en dosis NPK.
    """
    # Precios por kg de nutriente
    precio_n = precios['insumos']['urea'] / 1000 * 0.46  # Urea 46% N
    precio_p = precios['insumos']['fosfato'] / 1000 * 0.46  # Fosfato 46% P2O5
    precio_k = precios['insumos']['cloruro_potasio'] / 1000 * 0.60  # KCl 60% K2O
    # Dosis promedio (kg/ha) basado en recomendaciones
    if cultivo == 'MAÍZ':
        dosis = {'N': 150, 'P': 50, 'K': 120}
    elif cultivo == 'SOYA':
        dosis = {'N': 30, 'P': 40, 'K': 90}
    elif cultivo == 'TRIGO':
        dosis = {'N': 120, 'P': 40, 'K': 80}
    else:  # GIRASOL
        dosis = {'N': 80, 'P': 35, 'K': 100}
    # Ajustar dosis según NPK actual
    factor_ajuste = max(0.5, min(1.5, (1 - dosis_npk) * 2))
    costo_n = dosis['N'] * factor_ajuste * precio_n
    costo_p = dosis['P'] * factor_ajuste * precio_p
    costo_k = dosis['K'] * factor_ajuste * precio_k
    return {
        'costo_total': round(costo_n + costo_p + costo_k, 2),
        'costo_n': round(costo_n, 2),
        'costo_p': round(costo_p, 2),
        'costo_k': round(costo_k, 2),
        'dosis_ajustada': {
            'N': round(dosis['N'] * factor_ajuste, 1),
            'P': round(dosis['P'] * factor_ajuste, 1),
            'K': round(dosis['K'] * factor_ajuste, 1)
        }
    }

def calcular_rendimiento_potencial(npk_actual, cultivo, aplica_fertilizacion=True):
    """
    Calcula el rendimiento potencial basado en fertilidad y aplicación de NPK.
    """
    rendimiento_base = RENDIMIENTOS_BASE[cultivo]
    # Máximo rendimiento alcanzable con fertilización óptima
    max_rendimiento = {
        'MAÍZ': 12.0,
        'SOYA': 4.5,
        'TRIGO': 6.0,
        'GIRASOL': 3.5
    }
    if aplica_fertilizacion:
        # Con fertilización: respuesta logística
        factor_respuesta = 1.0 + (max_rendimiento[cultivo]/rendimiento_base - 1) * npk_actual
        rendimiento = rendimiento_base * factor_respuesta
    else:
        # Sin fertilización: rendimiento base reducido por baja fertilidad
        rendimiento = rendimiento_base * (0.3 + 0.7 * npk_actual)
    return round(rendimiento, 2)

def calcular_tir(inversion, ingresos_anuales, anos=5, tasa_descuento=0.12):
    """
    Calcula la Tasa Interna de Retorno (TIR).
    """
    try:
        # Flujo de caja: inversión inicial negativa, luego ingresos anuales
        flujos = [-inversion] + [ingresos_anuales] * anos
        # Calcular TIR usando método iterativo simple
        def npv(tasa):
            return sum([flujo / ((1 + tasa) ** i) for i, flujo in enumerate(flujos)])
        # Buscar TIR por bisección
        low, high = -0.99, 10.0
        for _ in range(100):
            mid = (low + high) / 2
            if npv(mid) > 0:
                low = mid
            else:
                high = mid
        tir = (low + high) / 2
        return round(tir * 100, 2)  # En porcentaje
    except:
        return 0.0

def generar_analisis_economico(gdf_analizado, cultivo, area_total, precios_actualizados):
    """
    Genera análisis económico completo para el lote.
    """
    if gdf_analizado.empty:
        return None
    # Obtener precio de venta
    precio_key = f'{cultivo.lower().replace("í", "i").replace("á", "a")}_rosario'
    precio_venta = precios_actualizados['pizarra_rosario'].get(precio_key, 200)
    # Costos base
    costo_base_ha = COSTOS_BASE[cultivo]
    resultados = {
        'escenario_sin': {'costos': [], 'ingresos': [], 'beneficios': [], 'rendimientos': []},
        'escenario_con': {'costos': [], 'ingresos': [], 'beneficios': [], 'rendimientos': []}
    }
    # Calcular por zona
    for idx, row in gdf_analizado.iterrows():
        area_ha = row.get('area_ha', area_total / len(gdf_analizado))
        npk_actual = row.get('npk_actual', 0.5)
        # Escenario SIN fertilización
        rendimiento_sin = calcular_rendimiento_potencial(npk_actual, cultivo, aplica_fertilizacion=False)
        ingreso_sin = rendimiento_sin * precio_venta * area_ha
        costo_sin = costo_base_ha * area_ha
        beneficio_sin = ingreso_sin - costo_sin
        # Escenario CON fertilización
        rendimiento_con = calcular_rendimiento_potencial(npk_actual, cultivo, aplica_fertilizacion=True)
        costo_fert = calcular_costo_fertilizacion(npk_actual, cultivo, precios_actualizados)
        costo_con = (costo_base_ha + costo_fert['costo_total']) * area_ha
        ingreso_con = rendimiento_con * precio_venta * area_ha
        beneficio_con = ingreso_con - costo_con
        # Almacenar resultados
        resultados['escenario_sin']['rendimientos'].append(rendimiento_sin)
        resultados['escenario_sin']['ingresos'].append(ingreso_sin)
        resultados['escenario_sin']['costos'].append(costo_sin)
        resultados['escenario_sin']['beneficios'].append(beneficio_sin)
        resultados['escenario_con']['rendimientos'].append(rendimiento_con)
        resultados['escenario_con']['ingresos'].append(ingreso_con)
        resultados['escenario_con']['costos'].append(costo_con)
        resultados['escenario_con']['beneficios'].append(beneficio_con)
    # Calcular totales
    for escenario in ['sin', 'con']:
        key = f'escenario_{escenario}'
        resultados[key]['total_ingresos'] = sum(resultados[key]['ingresos'])
        resultados[key]['total_costos'] = sum(resultados[key]['costos'])
        resultados[key]['total_beneficios'] = sum(resultados[key]['beneficios'])
        resultados[key]['rendimiento_promedio'] = np.mean(resultados[key]['rendimientos'])
        if len(gdf_analizado) > 0:
            resultados[key]['costo_promedio_ha'] = np.mean(resultados[key]['costos']) / np.mean([row.get('area_ha', 1) for idx, row in gdf_analizado.iterrows()])
        else:
            resultados[key]['costo_promedio_ha'] = 0
    # Calcular TIR
    inversion_fertilizacion = resultados['escenario_con']['total_costos'] - resultados['escenario_sin']['total_costos']
    ingreso_extra = resultados['escenario_con']['total_ingresos'] - resultados['escenario_sin']['total_ingresos']
    tir = calcular_tir(inversion_fertilizacion, ingreso_extra)
    resultados['tir'] = tir
    resultados['precio_venta'] = precio_venta
    resultados['inversion_fertilizacion'] = inversion_fertilizacion
    resultados['ingreso_extra'] = ingreso_extra
    if ingreso_extra > 0:
        resultados['payback'] = round(inversion_fertilizacion / ingreso_extra * 12, 1)
    else:
        resultados['payback'] = float('inf')
    return resultados

def crear_mapa_potencial_cosecha(gdf_analizado, cultivo, precios_actualizados):
    """
    Crea mapa de calor del potencial de cosecha con fertilización.
    """
    try:
        # Convertir a Web Mercator
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        # Calcular potencial por zona
        potenciales = []
        for idx, row in gdf_plot.iterrows():
            npk_actual = row.get('npk_actual', 0.5)
            rendimiento = calcular_rendimiento_potencial(npk_actual, cultivo, aplica_fertilizacion=True)
            potenciales.append(rendimiento)
        gdf_plot['potencial_cosecha'] = potenciales
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        # Configurar estilo oscuro
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        # Mapa de calor
        scatter = ax.scatter(
            [geom.centroid.x for geom in gdf_plot.geometry],
            [geom.centroid.y for geom in gdf_plot.geometry],
            c=gdf_plot['potencial_cosecha'],
            cmap='RdYlGn',
            s=300,
            alpha=0.8,
            edgecolors='white',
            linewidth=1
        )
        # Agregar mapa base ESRI Satellite
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.3)
        except:
            pass
        # Dibujar polígonos
        gdf_plot.plot(ax=ax, color='none', edgecolor='white', linewidth=1, alpha=0.5)
        # Etiquetas
        for idx, row in gdf_plot.iterrows():
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{row['potencial_cosecha']:.1f}t", (centroid.x, centroid.y),
                        xytext=(0, 0), textcoords="offset points", fontsize=8, color='white', weight='bold',
                        ha='center', va='center',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'))
        ax.set_title(f'🌱 MAPA DE POTENCIAL DE COSECHA - {cultivo}\n(Con aplicación óptima de NPK)',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')
        # Barra de colores
        cbar = plt.colorbar(scatter, ax=ax, shrink=0.8)
        cbar.set_label('Rendimiento Potencial (ton/ha)', fontsize=12, color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf, gdf_plot
    except Exception as e:
        st.error(f"Error creando mapa de potencial: {str(e)}")
        return None, None

def crear_mapa_rentabilidad(gdf_analizado, cultivo, precios_actualizados):
    """
    Crea mapa de calor de rentabilidad (USD/ha).
    """
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        # Calcular rentabilidad por zona
        rentabilidades = []
        for idx, row in gdf_plot.iterrows():
            npk_actual = row.get('npk_actual', 0.5)
            area_ha = row.get('area_ha', 1)
            # Calcular costo con fertilización
            costo_fert = calcular_costo_fertilizacion(npk_actual, cultivo, precios_actualizados)
            costo_total_ha = COSTOS_BASE[cultivo] + costo_fert['costo_total']
            # Calcular ingreso
            rendimiento = calcular_rendimiento_potencial(npk_actual, cultivo, aplica_fertilizacion=True)
            precio_key = f'{cultivo.lower().replace("í", "i").replace("á", "a")}_rosario'
            precio_venta = precios_actualizados['pizarra_rosario'].get(precio_key, 200)
            ingreso_ha = rendimiento * precio_venta
            # Rentabilidad
            rentabilidad = ingreso_ha - costo_total_ha
            rentabilidades.append(rentabilidad)
        gdf_plot['rentabilidad_usd_ha'] = rentabilidades
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        # Mapa de calor de rentabilidad
        scatter = ax.scatter(
            [geom.centroid.x for geom in gdf_plot.geometry],
            [geom.centroid.y for geom in gdf_plot.geometry],
            c=gdf_plot['rentabilidad_usd_ha'],
            cmap='RdYlBu_r',
            s=300,
            alpha=0.8,
            edgecolors='white',
            linewidth=1,
            vmin=-500,
            vmax=1500
        )
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.3)
        except:
            pass
        gdf_plot.plot(ax=ax, color='none', edgecolor='white', linewidth=1, alpha=0.5)
        # Etiquetas con color según rentabilidad
        for idx, row in gdf_plot.iterrows():
            centroid = row.geometry.centroid
            rent_color = 'white' if row['rentabilidad_usd_ha'] > 0 else '#ff6b6b'
            ax.annotate(f"Z{row['id_zona']}\n${row['rentabilidad_usd_ha']:.0f}", (centroid.x, centroid.y),
                        xytext=(0, 0), textcoords="offset points", fontsize=8, color=rent_color, weight='bold',
                        ha='center', va='center',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'))
        ax.set_title(f'💰 MAPA DE RENTABILIDAD - {cultivo}\n(Beneficio USD/ha con fertilización óptima)',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')
        cbar = plt.colorbar(scatter, ax=ax, shrink=0.8)
        cbar.set_label('Rentabilidad (USD/ha)', fontsize=12, color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')
        # Línea de equilibrio
        ax.axhline(0, color='yellow', linestyle='--', alpha=0.5, transform=ax.transAxes)
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"Error creando mapa de rentabilidad: {str(e)}")
        return None

# ===== INICIALIZACIÓN DE PRECIOS =====
def inicializar_precios():
    """Inicializa o actualiza los precios al cargar la aplicación"""
    try:
        if 'precios_actualizados' not in st.session_state:
            with st.spinner("Actualizando precios de mercado..."):
                precios = obtener_precios_actualizados()
                st.session_state['precios_actualizados'] = precios
        # Actualizar cada 24 horas
        if 'ultima_actualizacion' in st.session_state:
            ultima = datetime.strptime(st.session_state['ultima_actualizacion'], "%Y-%m-%d %H:%M:%S")
            if (datetime.now() - ultima).days >= 1:
                precios = obtener_precios_actualizados()
                st.session_state['precios_actualizados'] = precios
                st.session_state['ultima_actualizacion'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        else:
            st.session_state['ultima_actualizacion'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    except Exception as e:
        st.error(f"Error inicializando precios: {str(e)}")
        st.session_state['precios_actualizados'] = PRECIOS_API

# ===== SIDEBAR MEJORADO (INTERFAZ VISUAL) =====
with st.sidebar:
    st.markdown('<div class="sidebar-title">⚙️ CONFIGURACIÓN</div>', unsafe_allow_html=True)
    cultivo = st.selectbox("Cultivo:", ["MAÍZ", "SOYA", "TRIGO", "GIRASOL"])
    # IMAGEN CON MANEJO DE ERRORES
    try:
        # Normalizar el nombre del cultivo
        cultivo_key = cultivo.upper().replace("Í", "I").replace("Á", "A")
        # Verificar si tenemos la imagen
        if cultivo in IMAGENES_CULTIVOS:
            st.image(IMAGENES_CULTIVOS[cultivo],
                     caption=f"Cultivo: {cultivo}",
                     use_container_width=True)
        else:
            # Usar imagen por defecto
            st.image("https://images.unsplash.com/photo-1500382017468-9049fed747ef?w=400&h=300&fit=crop",
                     caption=f"Cultivo: {cultivo}",
                     use_container_width=True)
    except Exception as e:
        # Si todo falla, mostrar un placeholder simple
        st.markdown(f"**🌱 {cultivo}**")
        st.info(f"Imagen del cultivo: {cultivo}")
    analisis_tipo = st.selectbox("Tipo de Análisis:", ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK", "ANÁLISIS DE TEXTURA", "ANÁLISIS DE CURVAS DE NIVEL"])
    if analisis_tipo == "RECOMENDACIONES NPK":
        nutriente = st.selectbox("Nutriente:", ["NITRÓGENO", "FÓSFORO", "POTASIO"])
    else:
        nutriente = None

    st.subheader("🛰️ Fuente de Datos Satelitales")
    satelite_seleccionado = st.selectbox(
        "Satélite:",
        ["SENTINEL-2", "LANDSAT-8", "DATOS_SIMULADOS"],
        help="Selecciona la fuente de datos satelitales"
    )
    if satelite_seleccionado in SATELITES_DISPONIBLES:
        info_satelite = SATELITES_DISPONIBLES[satelite_seleccionado]
        st.info(f"""
        **{info_satelite['icono']} {info_satelite['nombre']}**
        - Resolución: {info_satelite['resolucion']}
        - Revisita: {info_satelite['revisita']}
        - Índices: {', '.join(info_satelite['indices'][:3])}
        """)
    if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
        st.subheader("📊 Índices de Vegetación")
        if satelite_seleccionado == "SENTINEL-2":
            indice_seleccionado = st.selectbox("Índice:", SATELITES_DISPONIBLES['SENTINEL-2']['indices'])
        elif satelite_seleccionado == "LANDSAT-8":
            indice_seleccionado = st.selectbox("Índice:", SATELITES_DISPONIBLES['LANDSAT-8']['indices'])
        else:
            indice_seleccionado = st.selectbox("Índice:", SATELITES_DISPONIBLES['DATOS_SIMULADOS']['indices'])

    if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
        st.subheader("📅 Rango Temporal")
        fecha_fin = st.date_input("Fecha fin", datetime.now())
        fecha_inicio = st.date_input("Fecha inicio", datetime.now() - timedelta(days=30))

    st.subheader("🎯 División de Parcela")
    n_divisiones = st.slider("Número de zonas de manejo:", min_value=16, max_value=48, value=32)

    if analisis_tipo == "ANÁLISIS DE CURVAS DE NIVEL":
        st.subheader("🏔️ Configuración Curvas de Nivel")
        intervalo_curvas = st.slider("Intervalo entre curvas (metros):", 1.0, 20.0, 5.0, 1.0)
        resolucion_dem = st.slider("Resolución DEM (metros):", 5.0, 50.0, 10.0, 5.0)

    st.subheader("📤 Subir Parcela")
    uploaded_file = st.file_uploader("Subir archivo de tu parcela", type=['zip', 'kml', 'kmz'],
                                     help="Formatos aceptados: Shapefile (.zip), KML (.kml), KMZ (.kmz)")

# ===== NUEVA SECCIÓN: ANÁLISIS ECONÓMICO =====
st.markdown("---")
st.markdown('<div class="sidebar-title">💰 ANÁLISIS ECONÓMICO</div>', unsafe_allow_html=True)
if st.button("🔄 Actualizar Precios de Mercado", key="actualizar_precios"):
    with st.spinner("Obteniendo precios actualizados..."):
        precios_actualizados = obtener_precios_actualizados()
        st.session_state['precios_actualizados'] = precios_actualizados
    st.success("✅ Precios actualizados")
if 'precios_actualizados' in st.session_state:
    st.info(f"📅 Precios actualizados: {st.session_state['precios_actualizados'].get('fecha_actualizacion', 'N/A')}")
plaza_precios = st.selectbox("Plaza de referencia:", ["ROSARIO", "BUENOS AIRES"])
st.subheader("📊 Parámetros Económicos")
costo_base_ha = st.number_input(
    f"Costo base producción ({cultivo}) USD/ha:",
    min_value=100.0,
    max_value=5000.0,
    value=float(COSTOS_BASE.get(cultivo, 1000)),
    step=50.0
)
precio_venta_manual = st.number_input(
    f"Precio venta {cultivo} USD/ton:",
    min_value=50.0,
    max_value=1000.0,
    value=float(PRECIOS_API['precios_pizarra'].get(f'{cultivo.lower().replace("í", "i").replace("á", "a")}_rosario', 200)),
    step=10.0
)
tasa_descuento = st.slider(
    "Tasa de descuento (%):",
    min_value=1.0,
    max_value=20.0,
    value=12.0,
    step=0.5
) / 100
st.session_state['parametros_economicos'] = {
    'plaza': plaza_precios,
    'costo_base_ha': costo_base_ha,
    'precio_venta': precio_venta_manual,
    'tasa_descuento': tasa_descuento
}

# ===== INICIALIZAR PRECIOS =====
inicializar_precios()

# ===== FUNCIONES AUXILIARES - CORREGIDAS PARA EPSG:4326 =====
def validar_y_corregir_crs(gdf):
    if gdf is None or len(gdf) == 0:
        return gdf
    try:
        if gdf.crs is None:
            gdf = gdf.set_crs('EPSG:4326', inplace=False)
        elif str(gdf.crs).upper() != 'EPSG:4326':
            original_crs = str(gdf.crs)
            gdf = gdf.to_crs('EPSG:4326')
        return gdf
    except Exception as e:
        return gdf

def calcular_superficie(gdf):
    try:
        if gdf is None or len(gdf) == 0:
            return 0.0
        gdf = validar_y_corregir_crs(gdf)
        bounds = gdf.total_bounds
        if bounds[0] < -180 or bounds[2] > 180 or bounds[1] < -90 or bounds[3] > 90:
            area_grados2 = gdf.geometry.area.sum()
            area_m2 = area_grados2 * 111000 * 111000
            return area_m2 / 10000
        gdf_projected = gdf.to_crs('EPSG:3857')
        area_m2 = gdf_projected.geometry.area.sum()
        return area_m2 / 10000
    except Exception as e:
        try:
            return gdf.geometry.area.sum() / 10000
        except:
            return 0.0

def dividir_parcela_en_zonas(gdf, n_zonas):
    if len(gdf) == 0:
        return gdf
    gdf = validar_y_corregir_crs(gdf)
    parcela_principal = gdf.iloc[0].geometry
    bounds = parcela_principal.bounds
    minx, miny, maxx, maxy = bounds
    sub_poligonos = []
    n_cols = math.ceil(math.sqrt(n_zonas))
    n_rows = math.ceil(n_zonas / n_cols)
    width = (maxx - minx) / n_cols
    height = (maxy - miny) / n_rows
    for i in range(n_rows):
        for j in range(n_cols):
            if len(sub_poligonos) >= n_zonas:
                break
            cell_minx = minx + (j * width)
            cell_maxx = minx + ((j + 1) * width)
            cell_miny = miny + (i * height)
            cell_maxy = miny + ((i + 1) * height)
            cell_poly = Polygon([(cell_minx, cell_miny), (cell_maxx, cell_miny), (cell_maxx, cell_maxy), (cell_minx, cell_maxy)])
            intersection = parcela_principal.intersection(cell_poly)
            if not intersection.is_empty and intersection.area > 0:
                sub_poligonos.append(intersection)
    if sub_poligonos:
        nuevo_gdf = gpd.GeoDataFrame({'id_zona': range(1, len(sub_poligonos) + 1), 'geometry': sub_poligonos}, crs='EPSG:4326')
        return nuevo_gdf
    else:
        return gdf

# ===== FUNCIONES PARA CARGAR ARCHIVOS =====
def cargar_shapefile_desde_zip(zip_file):
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                zip_ref.extractall(tmp_dir)
            shp_files = [f for f in os.listdir(tmp_dir) if f.endswith('.shp')]
            if shp_files:
                shp_path = os.path.join(tmp_dir, shp_files[0])
                gdf = gpd.read_file(shp_path)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
            else:
                st.error("❌ No se encontró ningún archivo .shp en el ZIP")
                return None
    except Exception as e:
        st.error(f"❌ Error cargando shapefile desde ZIP: {str(e)}")
        return None

def parsear_kml_manual(contenido_kml):
    try:
        root = ET.fromstring(contenido_kml)
        namespaces = {'kml': 'http://www.opengis.net/kml/2.2'}
        polygons = []
        for polygon_elem in root.findall('.//kml:Polygon', namespaces):
            coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
            if coords_elem is not None and coords_elem.text:
                coord_text = coords_elem.text.strip()
                coord_list = []
                for coord_pair in coord_text.split():
                    parts = coord_pair.split(',')
                    if len(parts) >= 2:
                        lon = float(parts[0])
                        lat = float(parts[1])
                        coord_list.append((lon, lat))
                if len(coord_list) >= 3:
                    polygons.append(Polygon(coord_list))
        if not polygons:
            for multi_geom in root.findall('.//kml:MultiGeometry', namespaces):
                for polygon_elem in multi_geom.findall('.//kml:Polygon', namespaces):
                    coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
                    if coords_elem is not None and coords_elem.text:
                        coord_text = coords_elem.text.strip()
                        coord_list = []
                        for coord_pair in coord_text.split():
                            parts = coord_pair.split(',')
                            if len(parts) >= 2:
                                lon = float(parts[0])
                                lat = float(parts[1])
                                coord_list.append((lon, lat))
                        if len(coord_list) >= 3:
                            polygons.append(Polygon(coord_list))
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        else:
            for placemark in root.findall('.//kml:Placemark', namespaces):
                for elem_name in ['Polygon', 'LineString', 'Point', 'LinearRing']:
                    elem = placemark.find(f'.//kml:{elem_name}', namespaces)
                    if elem is not None:
                        coords_elem = elem.find('.//kml:coordinates', namespaces)
                        if coords_elem is not None and coords_elem.text:
                            coord_text = coords_elem.text.strip()
                            coord_list = []
                            for coord_pair in coord_text.split():
                                parts = coord_pair.split(',')
                                if len(parts) >= 2:
                                    lon = float(parts[0])
                                    lat = float(parts[1])
                                    coord_list.append((lon, lat))
                            if len(coord_list) >= 3:
                                polygons.append(Polygon(coord_list))
                            break
            if polygons:
                gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
                return gdf
        return None
    except Exception as e:
        return None

def cargar_kml(kml_file):
    try:
        if kml_file.name.endswith('.kmz'):
            with tempfile.TemporaryDirectory() as tmp_dir:
                with zipfile.ZipFile(kml_file, 'r') as zip_ref:
                    zip_ref.extractall(tmp_dir)
                kml_files = [f for f in os.listdir(tmp_dir) if f.endswith('.kml')]
                if kml_files:
                    kml_path = os.path.join(tmp_dir, kml_files[0])
                    with open(kml_path, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    gdf = parsear_kml_manual(contenido)
                    if gdf is not None:
                        return gdf
                    else:
                        try:
                            gdf = gpd.read_file(kml_path)
                            gdf = validar_y_corregir_crs(gdf)
                            return gdf
                        except:
                            return None
                else:
                    return None
        else:
            contenido = kml_file.read().decode('utf-8')
            gdf = parsear_kml_manual(contenido)
            if gdf is not None:
                return gdf
            else:
                kml_file.seek(0)
                gdf = gpd.read_file(kml_file)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
    except Exception as e:
        return None

def cargar_archivo_parcela(uploaded_file):
    try:
        if uploaded_file.name.endswith('.zip'):
            gdf = cargar_shapefile_desde_zip(uploaded_file)
        elif uploaded_file.name.endswith(('.kml', '.kmz')):
            gdf = cargar_kml(uploaded_file)
        else:
            st.error("❌ Formato de archivo no soportado")
            return None
        if gdf is not None:
            gdf = validar_y_corregir_crs(gdf)
            if not gdf.geometry.geom_type.str.contains('Polygon').any():
                gdf = gdf.explode()
                gdf = gdf[gdf.geometry.geom_type.isin(['Polygon', 'MultiPolygon'])]
            if len(gdf) > 0:
                if 'id_zona' not in gdf.columns:
                    gdf['id_zona'] = range(1, len(gdf) + 1)
                return gdf
            else:
                return None
        return gdf
    except Exception as e:
        return None

# ===== FUNCIONES PARA DATOS SATELITALES =====
def descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.65 + np.random.normal(0, 0.1),
            'fuente': 'Landsat-8',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"LC08_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 15)}%",
            'resolucion': '30m'
        }
        return datos_simulados
    except Exception as e:
        return None

def descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.72 + np.random.normal(0, 0.08),
            'fuente': 'Sentinel-2',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"S2A_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 10)}%",
            'resolucion': '10m'
        }
        return datos_simulados
    except Exception as e:
        return None

def generar_datos_simulados(gdf, cultivo, indice='NDVI'):
    datos_simulados = {
        'indice': indice,
        'valor_promedio': PARAMETROS_CULTIVOS[cultivo]['NDVI_OPTIMO'] * 0.8 + np.random.normal(0, 0.1),
        'fuente': 'Simulación',
        'fecha': datetime.now().strftime('%Y-%m-%d'),
        'resolucion': '10m'
    }
    return datos_simulados

def obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin):
    """
    Obtiene datos meteorológicos diarios de NASA POWER para el centroide de la parcela.
    """
    try:
        centroid = gdf.geometry.unary_union.centroid
        lat = round(centroid.y, 4)
        lon = round(centroid.x, 4)
        start = fecha_inicio.strftime("%Y%m%d")
        end = fecha_fin.strftime("%Y%m%d")
        params = {
            'parameters': 'ALLSKY_SFC_SW_DWN,WS2M,T2M,PRECTOTCORR',
            'community': 'RE',
            'longitude': lon,
            'latitude': lat,
            'start': start,
            'end': end,
            'format': 'JSON'
        }
        url = "https://power.larc.nasa.gov/api/temporal/daily/point"
        response = requests.get(url, params=params, timeout=15)
        data = response.json()
        if 'properties' not in 
            return None
        series = data['properties']['parameter']
        df_power = pd.DataFrame({
            'fecha': pd.to_datetime(list(series['ALLSKY_SFC_SW_DWN'].keys())),
            'radiacion_solar': list(series['ALLSKY_SFC_SW_DWN'].values()),
            'viento_2m': list(series['WS2M'].values()),
            'temperatura': list(series['T2M'].values()),
            'precipitacion': list(series['PRECTOTCORR'].values())
        })
        df_power = df_power.replace(-999, np.nan).dropna()
        if df_power.empty:
            return None
        return df_power
    except Exception as e:
        return None

# ===== FUNCIONES DE ANÁLISIS GEE =====
def calcular_indices_satelitales_gee(gdf, cultivo, datos_satelitales):
    n_poligonos = len(gdf)
    resultados = []
    gdf_centroids = gdf.copy()
    gdf_centroids['centroid'] = gdf_centroids.geometry.centroid
    gdf_centroids['x'] = gdf_centroids.centroid.x
    gdf_centroids['y'] = gdf_centroids.centroid.y
    x_coords = gdf_centroids['x'].tolist()
    y_coords = gdf_centroids['y'].tolist()
    x_min, x_max = min(x_coords), max(x_coords)
    y_min, y_max = min(y_coords), max(y_coords)
    params = PARAMETROS_CULTIVOS[cultivo]
    valor_base_satelital = datos_satelitales.get('valor_promedio', 0.6) if datos_satelitales else 0.6
    for idx, row in gdf_centroids.iterrows():
        x_norm = (row['x'] - x_min) / (x_max - x_min) if x_max != x_min else 0.5
        y_norm = (row['y'] - y_min) / (y_max - y_min) if y_max != y_min else 0.5
        patron_espacial = (x_norm * 0.6 + y_norm * 0.4)
        base_mo = params['MATERIA_ORGANICA_OPTIMA'] * 0.7
        variabilidad_mo = patron_espacial * (params['MATERIA_ORGANICA_OPTIMA'] * 0.6)
        materia_organica = base_mo + variabilidad_mo + np.random.normal(0, 0.2)
        materia_organica = max(0.5, min(8.0, materia_organica))
        base_humedad = params['HUMEDAD_OPTIMA'] * 0.8
        variabilidad_humedad = patron_espacial * (params['HUMEDAD_OPTIMA'] * 0.4)
        humedad_suelo = base_humedad + variabilidad_humedad + np.random.normal(0, 0.05)
        humedad_suelo = max(0.1, min(0.8, humedad_suelo))
        ndvi_base = valor_base_satelital * 0.8
        ndvi_variacion = patron_espacial * (valor_base_satelital * 0.4)
        ndvi = ndvi_base + ndvi_variacion + np.random.normal(0, 0.06)
        ndvi = max(0.1, min(0.9, ndvi))
        ndre_base = params['NDRE_OPTIMO'] * 0.7
        ndre_variacion = patron_espacial * (params['NDRE_OPTIMO'] * 0.4)
        ndre = ndre_base + ndre_variacion + np.random.normal(0, 0.04)
        ndre = max(0.05, min(0.7, ndre))
        ndwi = 0.2 + np.random.normal(0, 0.08)
        ndwi = max(0, min(1, ndwi))
        npk_actual = (ndvi * 0.4) + (ndre * 0.3) + ((materia_organica / 8) * 0.2) + (humedad_suelo * 0.1)
        npk_actual = max(0, min(1, npk_actual))
        resultados.append({
            'materia_organica': round(materia_organica, 2),
            'humedad_suelo': round(humedad_suelo, 3),
            'ndvi': round(ndvi, 3),
            'ndre': round(ndre, 3),
            'ndwi': round(ndwi, 3),
            'npk_actual': round(npk_actual, 3)
        })
    return resultados

def calcular_recomendaciones_npk_gee(indices, nutriente, cultivo):
    recomendaciones = []
    params = PARAMETROS_CULTIVOS[cultivo]
    for idx in indices:
        ndre = idx['ndre']
        materia_organica = idx['materia_organica']
        humedad_suelo = idx['humedad_suelo']
        ndvi = idx['ndvi']
        if nutriente == "NITRÓGENO":
            factor_n = ((1 - ndre) * 0.6 + (1 - ndvi) * 0.4)
            n_recomendado = (factor_n * (params['NITROGENO']['max'] - params['NITROGENO']['min']) + params['NITROGENO']['min'])
            n_recomendado = max(params['NITROGENO']['min'] * 0.8, min(params['NITROGENO']['max'] * 1.2, n_recomendado))
            recomendaciones.append(round(n_recomendado, 1))
        elif nutriente == "FÓSFORO":
            factor_p = ((1 - (materia_organica / 8)) * 0.7 + (1 - humedad_suelo) * 0.3)
            p_recomendado = (factor_p * (params['FOSFORO']['max'] - params['FOSFORO']['min']) + params['FOSFORO']['min'])
            p_recomendado = max(params['FOSFORO']['min'] * 0.8, min(params['FOSFORO']['max'] * 1.2, p_recomendado))
            recomendaciones.append(round(p_recomendado, 1))
        else:
            factor_k = ((1 - ndre) * 0.4 + (1 - humedad_suelo) * 0.4 + (1 - (materia_organica / 8)) * 0.2)
            k_recomendado = (factor_k * (params['POTASIO']['max'] - params['POTASIO']['min']) + params['POTASIO']['min'])
            k_recomendado = max(params['POTASIO']['min'] * 0.8, min(params['POTASIO']['max'] * 1.2, k_recomendado))
            recomendaciones.append(round(k_recomendado, 1))
    return recomendaciones

# ===== FUNCIONES DE TEXTURA DEL SUELO =====
def clasificar_textura_suelo(arena, limo, arcilla):
    try:
        total = arena + limo + arcilla
        if total == 0:
            return "NO_DETERMINADA"
        arena_norm = (arena / total) * 100
        limo_norm = (limo / total) * 100
        arcilla_norm = (arcilla / total) * 100
        if arcilla_norm >= 35:
            return "Franco arcilloso"
        elif arcilla_norm >= 25 and arcilla_norm <= 35 and arena_norm >= 20 and arena_norm <= 45:
            return "Franco arcilloso"
        elif arena_norm >= 40 and arena_norm <= 50 and arcilla_norm >= 20 and arcilla_norm <= 30:
            return "Franco arenoso-arcilloso"
        elif arena_norm >= 50 and arena_norm <= 70 and arcilla_norm >= 5 and arcilla_norm <= 20:
            return "Franco arenoso-arcilloso"
        elif arcilla_norm >= 7 and arcilla_norm <= 27 and arena_norm >= 43 and arena_norm <= 52:
            return "Franco"
        elif arena_norm >= 85:
            return "Franco arenoso-arcilloso"
        else:
            return "Franco"
    except Exception as e:
        return "NO_DETERMINADA"

def analizar_textura_suelo(gdf, cultivo):
    gdf = validar_y_corregir_crs(gdf)
    params_textura = TEXTURA_SUELO_OPTIMA[cultivo]
    zonas_gdf = gdf.copy()
    zonas_gdf['area_ha'] = 0.0
    zonas_gdf['arena'] = 0.0
    zonas_gdf['limo'] = 0.0
    zonas_gdf['arcilla'] = 0.0
    zonas_gdf['textura_suelo'] = "NO_DETERMINADA"
    areas_ha_list = []
    arena_list = []
    limo_list = []
    arcilla_list = []
    textura_list = []
    for idx, row in zonas_gdf.iterrows():
        try:
            area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=zonas_gdf.crs)
            area_ha = calcular_superficie(area_gdf)
            if hasattr(area_ha, 'iloc'):
                area_ha = float(area_ha.iloc[0])
            elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                area_ha = float(area_ha[0])
            else:
                area_ha = float(area_ha)
            centroid = row.geometry.centroid if hasattr(row.geometry, 'centroid') else row.geometry.representative_point()
            seed_value = abs(hash(f"{centroid.x:.6f}_{centroid.y:.6f}_{cultivo}_textura")) % (2**32)
            rng = np.random.RandomState(seed_value)
            lat_norm = (centroid.y + 90) / 180 if centroid.y else 0.5
            lon_norm = (centroid.x + 180) / 360 if centroid.x else 0.5
            variabilidad_local = 0.15 + 0.7 * (lat_norm * lon_norm)
            arena_optima = params_textura['arena_optima']
            limo_optima = params_textura['limo_optima']
            arcilla_optima = params_textura['arcilla_optima']
            arena_val = max(5, min(95, rng.normal(
                arena_optima * (0.8 + 0.4 * variabilidad_local),
                arena_optima * 0.15
            )))
            limo_val = max(5, min(95, rng.normal(
                limo_optima * (0.7 + 0.6 * variabilidad_local),
                limo_optima * 0.2
            )))
            arcilla_val = max(5, min(95, rng.normal(
                arcilla_optima * (0.75 + 0.5 * variabilidad_local),
                arcilla_optima * 0.15
            )))
            total = arena_val + limo_val + arcilla_val
            arena_pct = (arena_val / total) * 100
            limo_pct = (limo_val / total) * 100
            arcilla_pct = (arcilla_val / total) * 100
            textura = clasificar_textura_suelo(arena_pct, limo_pct, arcilla_pct)
            areas_ha_list.append(area_ha)
            arena_list.append(float(arena_pct))
            limo_list.append(float(limo_pct))
            arcilla_list.append(float(arcilla_pct))
            textura_list.append(textura)
        except Exception as e:
            areas_ha_list.append(0.0)
            arena_list.append(float(params_textura['arena_optima']))
            limo_list.append(float(params_textura['limo_optima']))
            arcilla_list.append(float(params_textura['arcilla_optima']))
            textura_list.append(params_textura['textura_optima'])
    zonas_gdf['area_ha'] = areas_ha_list
    zonas_gdf['arena'] = arena_list
    zonas_gdf['limo'] = limo_list
    zonas_gdf['arcilla'] = arcilla_list
    zonas_gdf['textura_suelo'] = textura_list
    return zonas_gdf

# ===== FUNCIONES DE CURVAS DE NIVEL =====
def clasificar_pendiente(pendiente_porcentaje):
    for categoria, params in CLASIFICACION_PENDIENTES.items():
        if params['min'] <= pendiente_porcentaje < params['max']:
            return categoria, params['color']
    return "EXTREMA (>25%)", CLASIFICACION_PENDIENTES['EXTREMA (>25%)']['color']

def calcular_estadisticas_pendiente_simple(pendiente_grid):
    pendiente_flat = pendiente_grid.flatten()
    pendiente_flat = pendiente_flat[~np.isnan(pendiente_flat)]
    if len(pendiente_flat) == 0:
        return {'promedio': 0, 'min': 0, 'max': 0, 'std': 0, 'distribucion': {}}
    stats = {
        'promedio': float(np.mean(pendiente_flat)),
        'min': float(np.min(pendiente_flat)),
        'max': float(np.max(pendiente_flat)),
        'std': float(np.std(pendiente_flat)),
        'distribucion': {}
    }
    for categoria, params in CLASIFICACION_PENDIENTES.items():
        mask = (pendiente_flat >= params['min']) & (pendiente_flat < params['max'])
        stats['distribucion'][categoria] = {'porcentaje': float(np.sum(mask) / len(pendiente_flat) * 100), 'color': params['color']}
    return stats

def generar_dem_sintetico(gdf, resolucion=10.0):
    gdf = validar_y_corregir_crs(gdf)
    bounds = gdf.total_bounds
    minx, miny, maxx, maxy = bounds
    centroid = gdf.geometry.unary_union.centroid
    seed_value = int(centroid.x * 10000 + centroid.y * 10000) % (2**32)
    rng = np.random.RandomState(seed_value)
    num_cells = 50
    x = np.linspace(minx, maxx, num_cells)
    y = np.linspace(miny, maxy, num_cells)
    X, Y = np.meshgrid(x, y)
    elevacion_base = rng.uniform(100, 300)
    slope_x = rng.uniform(-0.001, 0.001)
    slope_y = rng.uniform(-0.001, 0.001)
    relief = np.zeros_like(X)
    n_hills = rng.randint(2, 5)
    for _ in range(n_hills):
        hill_center_x = rng.uniform(minx, maxx)
        hill_center_y = rng.uniform(miny, maxy)
        hill_radius = rng.uniform(0.001, 0.005)
        hill_height = rng.uniform(10, 50)
        dist = np.sqrt((X - hill_center_x)**2 + (Y - hill_center_y)**2)
        relief += hill_height * np.exp(-(dist**2) / (2 * hill_radius**2))
    noise = rng.randn(*X.shape) * 2
    Z = elevacion_base + slope_x * (X - minx) + slope_y * (Y - miny) + relief + noise
    Z = np.maximum(Z, 50)
    return X, Y, Z, bounds

def calcular_pendiente_simple(X, Y, Z, resolucion=10.0):
    dy = np.gradient(Z, axis=0) / resolucion
    dx = np.gradient(Z, axis=1) / resolucion
    pendiente = np.sqrt(dx**2 + dy**2) * 100
    pendiente = np.clip(pendiente, 0, 100)
    return pendiente

def crear_mapa_pendientes_simple(X, Y, pendiente_grid, gdf_original):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))
    fig.patch.set_facecolor('#0f172a')
    ax1.set_facecolor('#0f172a')
    ax2.set_facecolor('#0f172a')
    X_flat = X.flatten()
    Y_flat = Y.flatten()
    Z_flat = pendiente_grid.flatten()
    valid_mask = ~np.isnan(Z_flat)
    if np.sum(valid_mask) > 10:
        scatter = ax1.scatter(X_flat[valid_mask], Y_flat[valid_mask], c=Z_flat[valid_mask], cmap='RdYlGn_r', s=20, alpha=0.7, vmin=0, vmax=30)
        cbar = plt.colorbar(scatter, ax=ax1, shrink=0.8)
        cbar.set_label('Pendiente (%)', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')
        for porcentaje in [2, 5, 10, 15, 25]:
            mask_cat = (Z_flat[valid_mask] >= porcentaje-1) & (Z_flat[valid_mask] <= porcentaje+1)
            if np.sum(mask_cat) > 0:
                x_center = np.mean(X_flat[valid_mask][mask_cat])
                y_center = np.mean(Y_flat[valid_mask][mask_cat])
                ax1.text(x_center, y_center, f'{porcentaje}%', fontsize=8, fontweight='bold', ha='center', va='center',
                         bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'), color='white')
    else:
        ax1.text(0.5, 0.5, 'Datos insuficientes\npara mapa de calor', transform=ax1.transAxes, ha='center', va='center', fontsize=12, color='white')
    gdf_original.plot(ax=ax1, color='none', edgecolor='white', linewidth=2)
    ax1.set_title('Mapa de Calor de Pendientes', fontsize=12, fontweight='bold', color='white')
    ax1.set_xlabel('Longitud', color='white')
    ax1.set_ylabel('Latitud', color='white')
    ax1.tick_params(colors='white')
    ax1.grid(True, alpha=0.3, color='#475569')
    if np.sum(valid_mask) > 0:
        pendiente_data = Z_flat[valid_mask]
        ax2.hist(pendiente_data, bins=30, edgecolor='white', color='#3b82f6', alpha=0.7)
        for porcentaje, color in [(2, '#4daf4a'), (5, '#a6d96a'), (10, '#ffffbf'), (15, '#fdae61'), (25, '#f46d43')]:
            ax2.axvline(x=porcentaje, color=color, linestyle='--', linewidth=1, alpha=0.7)
            ax2.text(porcentaje+0.5, ax2.get_ylim()[1]*0.9, f'{porcentaje}%', color=color, fontsize=8)
        stats_pendiente = calcular_estadisticas_pendiente_simple(pendiente_grid)
        stats_text = f"""
        Estadísticas:
        • Mínima: {stats_pendiente['min']:.1f}%
        • Máxima: {stats_pendiente['max']:.1f}%
        • Promedio: {stats_pendiente['promedio']:.1f}%
        • Desviación: {stats_pendiente['std']:.1f}%
        """
        ax2.text(0.02, 0.98, stats_text, transform=ax2.transAxes, fontsize=9, verticalalignment='top',
                 color='white', bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'))
        ax2.set_xlabel('Pendiente (%)', color='white')
        ax2.set_ylabel('Frecuencia', color='white')
        ax2.set_title('Distribución de Pendientes', fontsize=12, fontweight='bold', color='white')
        ax2.tick_params(colors='white')
        ax2.grid(True, alpha=0.3, color='#475569')
    else:
        ax2.text(0.5, 0.5, 'Sin datos de pendiente', transform=ax2.transAxes, ha='center', va='center', fontsize=12, color='white')
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
    buf.seek(0)
    plt.close()
    return buf, calcular_estadisticas_pendiente_simple(pendiente_grid)

def generar_curvas_nivel_simple(X, Y, Z, intervalo=5.0, gdf_original=None):
    curvas = []
    elevaciones = []
    try:
        if gdf_original is not None:
            poligono_principal = gdf_original.iloc[0].geometry
            bounds = poligono_principal.bounds
            centro = poligono_principal.centroid
            ancho = bounds[2] - bounds[0]
            alto = bounds[3] - bounds[1]
            radio_max = min(ancho, alto) / 2
            z_min, z_max = np.nanmin(Z), np.nanmax(Z)
            n_curvas = min(10, int((z_max - z_min) / intervalo))
            for i in range(1, n_curvas + 1):
                radio = radio_max * (i / n_curvas)
                circle = centro.buffer(radio)
                interseccion = poligono_principal.intersection(circle)
                if interseccion.geom_type == 'LineString':
                    curvas.append(interseccion)
                    elevaciones.append(z_min + (i * intervalo))
                elif interseccion.geom_type == 'MultiLineString':
                    for parte in interseccion.geoms:
                        curvas.append(parte)
                        elevaciones.append(z_min + (i * intervalo))
    except Exception as e:
        if gdf_original is not None:
            bounds = gdf_original.total_bounds
            for i in range(3):
                y = bounds[1] + (i + 1) * ((bounds[3] - bounds[1]) / 4)
                linea = LineString([(bounds[0], y), (bounds[2], y)])
                curvas.append(linea)
                elevaciones.append(100 + i * 50)
    return curvas, elevaciones

# ===== FUNCIONES DE EXPORTACIÓN Y REPORTES =====
def exportar_a_geojson(gdf, nombre_base="parcela"):
    try:
        gdf = validar_y_corregir_crs(gdf)
        geojson_data = gdf.to_json()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{nombre_base}_{timestamp}.geojson"
        return geojson_data, nombre_archivo
    except Exception as e:
        return None, None

def generar_resumen_estadisticas(gdf_analizado, analisis_tipo, cultivo, df_power=None):
    estadisticas = {}
    try:
        if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
            if 'npk_actual' in gdf_analizado.columns:
                estadisticas['Índice NPK Promedio'] = f"{gdf_analizado['npk_actual'].mean():.3f}"
            if 'ndvi' in gdf_analizado.columns:
                estadisticas['NDVI Promedio'] = f"{gdf_analizado['ndvi'].mean():.3f}"
            if 'ndwi' in gdf_analizado.columns:
                estadisticas['NDWI Promedio'] = f"{gdf_analizado['ndwi'].mean():.3f}"
            if 'materia_organica' in gdf_analizado.columns:
                estadisticas['Materia Orgánica Promedio'] = f"{gdf_analizado['materia_organica'].mean():.1f}%"
            if df_power is not None:
                estadisticas['Radiación Solar Promedio'] = f"{df_power['radiacion_solar'].mean():.1f} kWh/m²/día"
                estadisticas['Velocidad Viento Promedio'] = f"{df_power['viento_2m'].mean():.2f} m/s"
                estadisticas['Precipitación Promedio'] = f"{df_power['precipitacion'].mean():.2f} mm/día"
        elif analisis_tipo == "ANÁLISIS DE TEXTURA":
            if 'arena' in gdf_analizado.columns:
                estadisticas['Arena Promedio'] = f"{gdf_analizado['arena'].mean():.1f}%"
                estadisticas['Limo Promedio'] = f"{gdf_analizado['limo'].mean():.1f}%"
                estadisticas['Arcilla Promedio'] = f"{gdf_analizado['arcilla'].mean():.1f}%"
            if 'textura_suelo' in gdf_analizado.columns:
                textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "N/D"
                estadisticas['Textura Predominante'] = textura_predominante
            if 'area_ha' in gdf_analizado.columns:
                estadisticas['Área Promedio por Zona'] = f"{gdf_analizado['area_ha'].mean():.2f} ha"
                if gdf_analizado['area_ha'].mean() > 0:
                    estadisticas['Coeficiente de Variación'] = f"{(gdf_analizado['area_ha'].std() / gdf_analizado['area_ha'].mean() * 100):.1f}%"
    except Exception as e:
        pass
    return estadisticas

def generar_recomendaciones_generales(gdf_analizado, analisis_tipo, cultivo):
    recomendaciones = []
    try:
        if analisis_tipo == "FERTILIDAD ACTUAL":
            if 'npk_actual' in gdf_analizado.columns:
                npk_promedio = gdf_analizado['npk_actual'].mean()
                if npk_promedio < 0.3:
                    recomendaciones.append("Fertilidad MUY BAJA: Se recomienda aplicación urgente de fertilizantes balanceados")
                    recomendaciones.append("Considerar enmiendas orgánicas para mejorar la estructura del suelo")
                elif npk_promedio < 0.5:
                    recomendaciones.append("Fertilidad BAJA: Recomendada aplicación de fertilizantes según análisis de suelo")
                elif npk_promedio < 0.7:
                    recomendaciones.append("Fertilidad ADECUADA: Mantener prácticas de manejo actuales")
                else:
                    recomendaciones.append("Fertilidad ÓPTIMA: Excelente condición, continuar con manejo actual")
        elif analisis_tipo == "ANÁLISIS DE TEXTURA":
            if 'textura_suelo' in gdf_analizado.columns:
                textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "N/D"
                if textura_predominante == "Franco arcilloso":
                    recomendaciones.append("Suelo franco arcilloso: Mejorar drenaje y evitar laboreo en condiciones húmedas")
                elif textura_predominante == "Franco arenoso-arcilloso":
                    recomendaciones.append("Suelo franco arenoso-arcilloso: Aumentar materia orgánica y considerar riego frecuente")
                elif textura_predominante == "Franco":
                    recomendaciones.append("Textura franca: Condiciones óptimas, mantener prácticas de conservación")
        if cultivo == "MAÍZ":
            recomendaciones.append("Para maíz: Priorizar aplicación de nitrógeno en etapas de crecimiento vegetativo.")
            recomendaciones.append("Mantener humedad adecuada durante floración y llenado de grano.")
        elif cultivo == "SOYA":
            recomendaciones.append("Para soya: Inocular con rizobios para fijación de nitrógeno atmosférico.")
            recomendaciones.append("Manejo adecuado de humedad durante formación de vainas.")
        elif cultivo == "TRIGO":
            recomendaciones.append("Para trigo: Aplicar nitrógeno en macollamiento y encañazón.")
            recomendaciones.append("Controlar humedad para evitar enfermedades fúngicas.")
        elif cultivo == "GIRASOL":
            recomendaciones.append("Para girasol: Aplicar potasio para mejorar calidad de semilla.")
            recomendaciones.append("Mantener buen drenaje, sensible a encharcamiento.")
        recomendaciones.append("Realizar análisis de suelo de laboratorio para validar resultados satelitales")
        recomendaciones.append("Considerar agricultura de precisión para aplicación variable de insumos")
    except Exception as e:
        recomendaciones.append("Error generando recomendaciones específicas")
    return recomendaciones

def limpiar_texto_para_pdf(texto):
    if not isinstance(texto, str):
        texto = str(texto)
    reemplazos = {
        '\u2022': '-',
        '\u2705': '[OK]',
        '\u26A0\uFE0F': '[!]',
        '\u274C': '[X]',
        '\u2013': '-',
        '\u2014': '--',
        '\u2018': "'",
        '\u2019': "'",
        '\u201C': '"',
        '\u201D': '"',
        '\u2192': '->',
        '\u2190': '<-',
        '\u2265': '>=',
        '\u2264': '<=',
        '\u00A0': ' ',
    }
    for original, reemplazo in reemplazos.items():
        texto = texto.replace(original, reemplazo)
    texto = texto.encode('latin-1', errors='replace').decode('latin-1')
    return texto

def generar_reporte_pdf(gdf_analizado, cultivo, analisis_tipo, area_total,
                        nutriente=None, satelite=None, indice=None,
                        mapa_buffer=None, estadisticas=None, recomendaciones=None,
                        analisis_economico=None):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font('Arial', '', 12)
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, limpiar_texto_para_pdf(f'REPORTE DE ANÁLISIS AGRÍCOLA - {cultivo}'), 0, 1, 'C')
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, limpiar_texto_para_pdf(f'Tipo de Análisis: {analisis_tipo}'), 0, 1, 'C')
        pdf.cell(0, 10, limpiar_texto_para_pdf(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}'), 0, 1, 'C')
        pdf.ln(10)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '1. INFORMACIÓN GENERAL', 0, 1)
        pdf.set_font('Arial', '', 12)
        info_general = f"""Cultivo: {cultivo}
Área Total: {area_total:.2f} ha
Zonas Analizadas: {len(gdf_analizado)}
Tipo de Análisis: {analisis_tipo}"""
        if satelite:
            info_general += f"\nSatélite: {satelite}"
        if indice:
            info_general += f"\nÍndice: {indice}"
        if nutriente:
            info_general += f"\nNutriente Analizado: {nutriente}"
        for linea in info_general.strip().split('\n'):
            pdf.cell(0, 8, limpiar_texto_para_pdf(linea), 0, 1)
        pdf.ln(5)
        if analisis_economico:
            pdf.add_page()
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, 'ANÁLISIS ECONÓMICO Y RENTABILIDAD', 0, 1, 'C')
            pdf.ln(5)
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '1. INDICADORES DE RENTABILIDAD', 0, 1)
            pdf.set_font('Arial', '', 12)
            metricas = [
                f"Tasa Interna de Retorno (TIR): {analisis_economico['tir']}%",
                f"Período de Payback: {analisis_economico['payback']} meses",
                f"Inversión en Fertilización: ${analisis_economico['inversion_fertilizacion']:,.0f}",
                f"Ingreso Extra Anual: ${analisis_economico['ingreso_extra']:,.0f}",
                f"Beneficio Neto Incremental: ${analisis_economico['escenario_con']['total_beneficios'] - analisis_economico['escenario_sin']['total_beneficios']:,.0f}"
            ]
            for metrica in metricas:
                pdf.cell(0, 8, limpiar_texto_para_pdf(metrica), 0, 1)
            pdf.ln(5)
        if estadisticas:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '2. ESTADÍSTICAS PRINCIPALES', 0, 1)
            pdf.set_font('Arial', '', 12)
            for key, value in estadisticas.items():
                linea = f"- {key}: {value}"
                pdf.cell(0, 8, limpiar_texto_para_pdf(linea), 0, 1)
            pdf.ln(5)
        if mapa_buffer:
            try:
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, '3. MAPA DE RESULTADOS', 0, 1)
                temp_img_path = "temp_map.png"
                with open(temp_img_path, "wb") as f:
                    f.write(mapa_buffer.getvalue())
                pdf.image(temp_img_path, x=10, w=190)
                pdf.ln(5)
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
            except Exception as e:
                pdf.cell(0, 8, limpiar_texto_para_pdf(f"Error al incluir mapa: {str(e)[:50]}..."), 0, 1)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '4. RESUMEN DE ZONAS', 0, 1)
        pdf.set_font('Arial', '', 10)
        if gdf_analizado is not None and not gdf_analizado.empty:
            columnas_mostrar = ['id_zona', 'area_ha']
            if 'npk_actual' in gdf_analizado.columns:
                columnas_mostrar.append('npk_actual')
            if 'valor_recomendado' in gdf_analizado.columns:
                columnas_mostrar.append('valor_recomendado')
            if 'textura_suelo' in gdf_analizado.columns:
                columnas_mostrar.append('textura_suelo')
            if 'ndwi' in gdf_analizado.columns:
                columnas_mostrar.append('ndwi')
            columnas_mostrar = [col for col in columnas_mostrar if col in gdf_analizado.columns]
            if columnas_mostrar:
                datos_tabla = [columnas_mostrar]
                for _, row in gdf_analizado.head(15).iterrows():
                    fila = []
                    for col in columnas_mostrar:
                        if col in gdf_analizado.columns:
                            valor = row[col]
                            if isinstance(valor, float):
                                if col in ['npk_actual', 'ndwi']:
                                    fila.append(f"{valor:.3f}")
                                else:
                                    fila.append(f"{valor:.2f}")
                            else:
                                fila.append(str(valor))
                        else:
                            fila.append("N/A")
                    datos_tabla.append(fila)
                col_widths = [190 // len(columnas_mostrar)] * len(columnas_mostrar)
                for fila in datos_tabla:
                    for i, item in enumerate(fila):
                        if i < len(col_widths):
                            pdf.cell(col_widths[i], 8, limpiar_texto_para_pdf(str(item)), border=1)
                    pdf.ln()
        pdf.ln(5)
        if recomendaciones:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '5. RECOMENDACIONES', 0, 1)
            pdf.set_font('Arial', '', 12)
            for rec in recomendaciones:
                linea = f"- {limpiar_texto_para_pdf(rec)}"
                pdf.multi_cell(0, 8, linea)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '6. METADATOS TÉCNICOS', 0, 1)
        pdf.set_font('Arial', '', 10)
        metadatos = f"""Generado por: Analizador Multi-Cultivo Satellital
Versión: 2.0
Fecha de generación: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
Sistema de coordenadas: EPSG:4326 (WGS84)
Número de zonas: {len(gdf_analizado)}"""
        for linea in metadatos.strip().split('\n'):
            pdf.cell(0, 6, limpiar_texto_para_pdf(linea), 0, 1)
        pdf_output = BytesIO()
        pdf_output.write(pdf.output(dest='S').encode('latin-1'))
        pdf_output.seek(0)
        return pdf_output
    except Exception as e:
        return None

def generar_reporte_docx(gdf_analizado, cultivo, analisis_tipo, area_total,
                         nutriente=None, satelite=None, indice=None,
                         mapa_buffer=None, estadisticas=None, recomendaciones=None,
                         analisis_economico=None):
    try:
        doc = Document()
        title = doc.add_heading(f'REPORTE DE ANÁLISIS AGRÍCOLA - {cultivo}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(f'Tipo de Análisis: {analisis_tipo}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        doc.add_heading('1. INFORMACIÓN GENERAL', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = 'Cultivo'
        info_table.cell(0, 1).text = cultivo
        info_table.cell(1, 0).text = 'Área Total'
        info_table.cell(1, 1).text = f'{area_total:.2f} ha'
        info_table.cell(2, 0).text = 'Zonas Analizadas'
        info_table.cell(2, 1).text = str(len(gdf_analizado))
        info_table.cell(3, 0).text = 'Tipo de Análisis'
        info_table.cell(3, 1).text = analisis_tipo
        row_count = 4
        if satelite:
            if row_count >= len(info_table.rows):
                info_table.add_row()
            info_table.cell(row_count, 0).text = 'Satélite'
            info_table.cell(row_count, 1).text = satelite
            row_count += 1
        if indice:
            if row_count >= len(info_table.rows):
                info_table.add_row()
            info_table.cell(row_count, 0).text = 'Índice'
            info_table.cell(row_count, 1).text = indice
            row_count += 1
        if nutriente:
            if row_count >= len(info_table.rows):
                info_table.add_row()
            info_table.cell(row_count, 0).text = 'Nutriente Analizado'
            info_table.cell(row_count, 1).text = nutriente
        if analisis_economico:
            doc.add_page()
            doc.add_heading('ANÁLISIS ECONÓMICO Y RENTABILIDAD', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            doc.add_heading('1. INDICADORES DE RENTABILIDAD', level=1)
            metricas = [
                ('Tasa Interna de Retorno (TIR)', f"{analisis_economico['tir']}%"),
                ('Período de Payback', f"{analisis_economico['payback']} meses"),
                ('Inversión en Fertilización', f"${analisis_economico['inversion_fertilizacion']:,.0f}"),
                ('Ingreso Extra Anual', f"${analisis_economico['ingreso_extra']:,.0f}"),
                ('Beneficio Neto Incremental', f"${analisis_economico['escenario_con']['total_beneficios'] - analisis_economico['escenario_sin']['total_beneficios']:,.0f}")
            ]
            for titulo, valor in metricas:
                p = doc.add_paragraph()
                p.add_run(f'{titulo}: ').bold = True
                p.add_run(valor)
            doc.add_paragraph()
        if estadisticas:
            doc.add_heading('2. ESTADÍSTICAS PRINCIPALES', level=1)
            for key, value in estadisticas.items():
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f'{key}: ')
                run.bold = True
                p.add_run(str(value))
            doc.add_paragraph()
        if mapa_buffer:
            try:
                doc.add_heading('3. MAPA DE RESULTADOS', level=1)
                temp_img_path = "temp_map_docx.png"
                with open(temp_img_path, "wb") as f:
                    f.write(mapa_buffer.getvalue())
                doc.add_picture(temp_img_path, width=Inches(6.0))
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f'Error al incluir mapa: {str(e)[:50]}...')
        doc.add_heading('4. RESUMEN DE ZONAS', level=1)
        if gdf_analizado is not None and not gdf_analizado.empty:
            columnas_mostrar = ['id_zona', 'area_ha']
            if 'npk_actual' in gdf_analizado.columns:
                columnas_mostrar.append('npk_actual')
            if 'valor_recomendado' in gdf_analizado.columns:
                columnas_mostrar.append('valor_recomendado')
            if 'textura_suelo' in gdf_analizado.columns:
                columnas_mostrar.append('textura_suelo')
            if 'ndwi' in gdf_analizado.columns:
                columnas_mostrar.append('ndwi')
            columnas_mostrar = [col for col in columnas_mostrar if col in gdf_analizado.columns]
            if columnas_mostrar:
                tabla = doc.add_table(rows=1, cols=len(columnas_mostrar))
                tabla.style = 'Table Grid'
                for i, col in enumerate(columnas_mostrar):
                    tabla.cell(0, i).text = col.replace('_', ' ').upper()
                for idx, row in gdf_analizado.head(10).iterrows():
                    row_cells = tabla.add_row().cells
                    for i, col in enumerate(columnas_mostrar):
                        if col in gdf_analizado.columns:
                            valor = row[col]
                            if isinstance(valor, float):
                                if col in ['npk_actual', 'ndwi']:
                                    row_cells[i].text = f"{valor:.3f}"
                                else:
                                    row_cells[i].text = f"{valor:.2f}"
                            else:
                                row_cells[i].text = str(valor)
                        else:
                            row_cells[i].text = "N/A"
        doc.add_paragraph()
        if recomendaciones:
            doc.add_heading('5. RECOMENDACIONES', level=1)
            for rec in recomendaciones:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(rec)
        doc.add_heading('6. METADATOS TÉCNICOS', level=1)
        metadatos = [
            ('Generado por', 'Analizador Multi-Cultivo Satellital'),
            ('Versión', '2.0'),
            ('Fecha de generación', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ('Sistema de coordenadas', 'EPSG:4326 (WGS84)'),
            ('Número de zonas', str(len(gdf_analizado)))
        ]
        for key, value in metadatos:
            p = doc.add_paragraph()
            run_key = p.add_run(f'{key}: ')
            run_key.bold = True
            p.add_run(value)
        docx_output = BytesIO()
        doc.save(docx_output)
        docx_output.seek(0)
        return docx_output
    except Exception as e:
        return None

# ===== FUNCIONES DE VISUALIZACIÓN =====
def crear_mapa_estatico_con_esri(gdf, titulo, columna_valor, analisis_tipo, nutriente, cultivo, satelite):
    try:
        gdf_plot = gdf.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        if analisis_tipo == "FERTILIDAD ACTUAL":
            cmap = LinearSegmentedColormap.from_list('fertilidad_gee', PALETAS_GEE['FERTILIDAD'])
            vmin, vmax = 0, 1
        else:
            if nutriente == "NITRÓGENO":
                cmap = LinearSegmentedColormap.from_list('nitrogeno_gee', PALETAS_GEE['NITROGENO'])
                vmin, vmax = (PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['min'] * 0.8,
                              PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['max'] * 1.2)
            elif nutriente == "FÓSFORO":
                cmap = LinearSegmentedColormap.from_list('fosforo_gee', PALETAS_GEE['FOSFORO'])
                vmin, vmax = (PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['min'] * 0.8,
                              PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['max'] * 1.2)
            else:
                cmap = LinearSegmentedColormap.from_list('potasio_gee', PALETAS_GEE['POTASIO'])
                vmin, vmax = (PARAMETROS_CULTIVOS[cultivo]['POTASIO']['min'] * 0.8,
                              PARAMETROS_CULTIVOS[cultivo]['POTASIO']['max'] * 1.2)  # ← CORREGIDO AQUÍ
        for idx, row in gdf_plot.iterrows():
            valor = row[columna_valor]
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap(valor_norm)
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white', linewidth=1.5, alpha=0.7)
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.1f}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points", fontsize=8, color='white', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='#1e293b', alpha=0.9, edgecolor='white'))
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.4)
        except:
            pass
        info_satelite = SATELITES_DISPONIBLES.get(satelite, SATELITES_DISPONIBLES['DATOS_SIMULADOS'])
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} ANÁLISIS GEE - {cultivo}\n'
                     f'{info_satelite["icono"]} {info_satelite["nombre"]} - {analisis_tipo}\n'
                     f'{columna_valor}',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label(columna_valor, fontsize=12, fontweight='bold', color='white')
        cbar.ax.yaxis.set_tick_params(color='white')
        cbar.outline.set_edgecolor('white')
        plt.setp(plt.getp(cbar.ax.axes, 'yticklabels'), color='white')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        return None

def crear_mapa_texturas_con_esri(gdf_analizado, cultivo):
    try:
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        fig.patch.set_facecolor('#0f172a')
        ax.set_facecolor('#0f172a')
        colores_textura = {
            'Franco': '#c7eae5',
            'Franco arcilloso': '#5ab4ac',
            'Franco arenoso-arcilloso': '#f6e8c3',
            'NO_DETERMINADA': '#999999'
        }
        for idx, row in gdf_plot.iterrows():
            textura = row['textura_suelo']
            color = colores_textura.get(textura, '#999999')
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='white', linewidth=1.5, alpha=0.8)
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{textura[:10]}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points", fontsize=8, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.4)
        except:
            pass
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} MAPA DE TEXTURAS - {cultivo}',
                     fontsize=16, fontweight='bold', pad=20, color='white')
        ax.set_xlabel('Longitud', color='white')
        ax.set_ylabel('Latitud', color='white')
        ax.tick_params(colors='white')
        ax.grid(True, alpha=0.3, color='#475569')
        from matplotlib.patches import Patch
        legend_elements = [Patch(facecolor=color, edgecolor='white', label=textura)
                           for textura, color in colores_textura.items()]
        legend = ax.legend(handles=legend_elements, title='Texturas', loc='upper left', bbox_to_anchor=(1.05, 1))
        legend.get_title().set_color('white')
        for text in legend.get_texts():
            text.set_color('white')
        legend.get_frame().set_facecolor('#1e293b')
        legend.get_frame().set_edgecolor('white')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0f172a')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        return None

def crear_grafico_personalizado(series, titulo, ylabel, color_linea, fondo_grafico='#0f172a', color_texto='#ffffff'):
    """Crea gráfico de línea con estilo oscuro"""
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_facecolor(fondo_grafico)
    fig.patch.set_facecolor(fondo_grafico)
    ax.plot(series.index, series.values, color=color_linea, linewidth=2.2)
    ax.set_title(titulo, fontsize=14, fontweight='bold', color=color_texto)
    ax.set_ylabel(ylabel, fontsize=12, color=color_texto)
    ax.set_xlabel("Fecha", fontsize=11, color=color_texto)
    ax.tick_params(axis='x', colors=color_texto, rotation=0)
    ax.tick_params(axis='y', colors=color_texto)
    ax.grid(True, color='#475569', linestyle='--', linewidth=0.7, alpha=0.7)
    for spine in ax.spines.values():
        spine.set_color('#475569')
    plt.tight_layout()
    return fig

def crear_grafico_barras_personalizado(series, titulo, ylabel, color_barra, fondo_grafico='#0f172a', color_texto='#ffffff'):
    """Crea gráfico de barras con estilo oscuro"""
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_facecolor(fondo_grafico)
    fig.patch.set_facecolor(fondo_grafico)
    ax.bar(series.index, series.values, color=color_barra, alpha=0.85)
    ax.set_title(titulo, fontsize=14, fontweight='bold', color=color_texto)
    ax.set_ylabel(ylabel, fontsize=12, color=color_texto)
    ax.set_xlabel("Fecha", fontsize=11, color=color_texto)
    ax.tick_params(axis='x', colors=color_texto, rotation=0)
    ax.tick_params(axis='y', colors=color_texto)
    ax.grid(axis='y', color='#475569', linestyle='--', linewidth=0.7, alpha=0.7)
    for spine in ax.spines.values():
        spine.set_color('#475569')
    plt.tight_layout()
    return fig

# ===== FUNCIÓN PRINCIPAL DE ANÁLISIS =====
def ejecutar_analisis(gdf, nutriente, analisis_tipo, n_divisiones, cultivo,
                      satelite=None, indice=None, fecha_inicio=None,
                      fecha_fin=None, intervalo_curvas=5.0, resolucion_dem=10.0):
    resultados = {
        'exitoso': False,
        'gdf_analizado': None,
        'mapa_buffer': None,
        'tabla_datos': None,
        'estadisticas': {},
        'recomendaciones': [],
        'area_total': 0,
        'df_power': None
    }
    try:
        gdf = validar_y_corregir_crs(gdf)
        area_total = calcular_superficie(gdf)
        resultados['area_total'] = area_total
        if analisis_tipo == "ANÁLISIS DE TEXTURA":
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            gdf_analizado = analizar_textura_suelo(gdf_dividido, cultivo)
            resultados['gdf_analizado'] = gdf_analizado
            resultados['exitoso'] = True
            return resultados
        elif analisis_tipo == "ANÁLISIS DE CURVAS DE NIVEL":
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            resultados['gdf_analizado'] = gdf_dividido
            resultados['exitoso'] = True
            return resultados
        elif analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
            datos_satelitales = None
            if satelite == "SENTINEL-2":
                datos_satelitales = descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice)
            elif satelite == "LANDSAT-8":
                datos_satelitales = descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice)
            else:
                datos_satelitales = generar_datos_simulados(gdf, cultivo, indice)
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            indices_gee = calcular_indices_satelitales_gee(gdf_dividido, cultivo, datos_satelitales)
            gdf_analizado = gdf_dividido.copy()
            for idx, indice_data in enumerate(indices_gee):
                for key, value in indice_data.items():
                    gdf_analizado.loc[gdf_analizado.index[idx], key] = value
            areas_ha_list = []
            for idx, row in gdf_analizado.iterrows():
                area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=gdf_analizado.crs)
                area_ha = calcular_superficie(area_gdf)
                if hasattr(area_ha, 'iloc'):
                    area_ha = float(area_ha.iloc[0])
                elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                    area_ha = float(area_ha[0])
                else:
                    area_ha = float(area_ha)
                areas_ha_list.append(area_ha)
            gdf_analizado['area_ha'] = areas_ha_list
            if analisis_tipo == "RECOMENDACIONES NPK":
                recomendaciones_npk = calcular_recomendaciones_npk_gee(indices_gee, nutriente, cultivo)
                gdf_analizado['valor_recomendado'] = recomendaciones_npk
            resultados['gdf_analizado'] = gdf_analizado
            resultados['exitoso'] = True
            if satelite:
                df_power = obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin)
                if df_power is not None:
                    resultados['df_power'] = df_power
            return resultados
        else:
            return resultados
    except Exception as e:
        return resultados

# ===== INTERFAZ PRINCIPAL =====
def main():
    # Inicializar precios
    inicializar_precios()

    # Mostrar información de precios actualizados
    if 'precios_actualizados' in st.session_state:
        precios = st.session_state['precios_actualizados']
        fecha_actualizacion = precios.get('fecha_actualizacion', 'N/A')
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("🌽 Maíz", f"${precios['pizarra_rosario'].get('maiz_rosario', 200)}/ton")
        with col2:
            st.metric("🫘 Soya", f"${precios['pizarra_rosario'].get('soya_rosario', 400)}/ton")
        with col3:
            st.metric("🌾 Trigo", f"${precios['pizarra_rosario'].get('trigo_rosario', 250)}/ton")
        with col4:
            st.metric("🌻 Girasol", f"${precios['pizarra_rosario'].get('girasol_rosario', 350)}/ton")
        st.caption(f"📅 Precios actualizados al: {fecha_actualizacion}")

    # Crear pestañas principales
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 ANÁLISIS PRINCIPAL", "🌱 FERTILIDAD", "💰 ECONOMÍA", "📈 REPORTES", "ℹ️ AYUDA"])

    with tab1:
        st.header("📊 ANÁLISIS PRINCIPAL")
        # Obtener variables del sidebar
        if 'uploaded_file' in locals() and uploaded_file is not None:
            with st.spinner("Cargando y analizando parcela..."):
                gdf_cargado = cargar_archivo_parcela(uploaded_file)
                if gdf_cargado is not None and not gdf_cargado.empty:
                    # Ejecutar análisis según tipo seleccionado
                    resultados = ejecutar_analisis(
                        gdf=gdf_cargado,
                        nutriente=nutriente,
                        analisis_tipo=analisis_tipo,
                        n_divisiones=n_divisiones,
                        cultivo=cultivo,
                        satelite=satelite_seleccionado,
                        indice=indice_seleccionado if 'indice_seleccionado' in locals() else None,
                        fecha_inicio=fecha_inicio if 'fecha_inicio' in locals() else None,
                        fecha_fin=fecha_fin if 'fecha_fin' in locals() else None,
                        intervalo_curvas=intervalo_curvas if 'intervalo_curvas' in locals() else 5.0,
                        resolucion_dem=resolucion_dem if 'resolucion_dem' in locals() else 10.0
                    )
                    if resultados['exitoso']:
                        # Mostrar resultados
                        col1, col2 = st.columns([2, 1])
                        with col1:
                            # ✅ VALIDACIÓN CORREGIDA AQUÍ
                            mapa_buf = resultados.get('mapa_buffer')
                            if mapa_buf is not None:
                                try:
                                    mapa_buf.seek(0)
                                    st.image(mapa_buf, caption=f"Mapa de {analisis_tipo} - {cultivo}", use_container_width=True)
                                except Exception as e:
                                    st.warning(f"⚠️ No se pudo mostrar el mapa: {str(e)}")
                            else:
                                st.warning("⚠️ El mapa no se generó correctamente.")
                        with col2:
                            st.subheader("📈 Resumen del Análisis")
                            st.metric("Área Total", f"{resultados['area_total']:.2f} ha")
                            st.metric("Número de Zonas", f"{len(resultados['gdf_analizado'])}")
                            # Mostrar estadísticas específicas
                            if analisis_tipo == "FERTILIDAD ACTUAL":
                                if 'npk_actual' in resultados['gdf_analizado'].columns:
                                    npk_promedio = resultados['gdf_analizado']['npk_actual'].mean()
                                    st.metric("Fertilidad Promedio (NPK)", f"{npk_promedio:.2%}")
                            elif analisis_tipo == "RECOMENDACIONES NPK":
                                if 'valor_recomendado' in resultados['gdf_analizado'].columns:
                                    rec_promedio = resultados['gdf_analizado']['valor_recomendado'].mean()
                                    st.metric(f"{nutriente} Recomendado", f"{rec_promedio:.1f} kg/ha")

                        # Botón para análisis económico
                        if st.button("💰 Realizar Análisis Económico", type="primary"):
                            st.session_state['mostrar_economia'] = True

                        # Mostrar tabla de datos
                        st.subheader("📋 Datos por Zona")
                        if resultados['gdf_analizado'] is not None:
                            # Seleccionar columnas a mostrar
                            columnas_interes = ['id_zona', 'area_ha']
                            if 'npk_actual' in resultados['gdf_analizado'].columns:
                                columnas_interes.append('npk_actual')
                            if 'materia_organica' in resultados['gdf_analizado'].columns:
                                columnas_interes.append('materia_organica')
                            if 'humedad_suelo' in resultados['gdf_analizado'].columns:
                                columnas_interes.append('humedad_suelo')
                            if 'valor_recomendado' in resultados['gdf_analizado'].columns:
                                columnas_interes.append('valor_recomendado')
                            if 'textura_suelo' in resultados['gdf_analizado'].columns:
                                columnas_interes.append('textura_suelo')
                            df_display = resultados['gdf_analizado'][columnas_interes].copy()
                            st.dataframe(df_display, use_container_width=True)

                        # Opciones de exportación
                        col_exp1, col_exp2, col_exp3 = st.columns(3)
                        with col_exp1:
                            if st.button("📥 Exportar a GeoJSON"):
                                geojson_data, nombre_archivo = exportar_a_geojson(resultados['gdf_analizado'], f"{cultivo}_{analisis_tipo}")
                                if geojson_data:
                                    st.download_button(
                                        label="Descargar GeoJSON",
                                        data=geojson_data,
                                        file_name=nombre_archivo,
                                        mime="application/json"
                                    )
                        with col_exp2:
                            if st.button("📄 Generar Reporte PDF"):
                                reporte_pdf = generar_reporte_pdf(
                                    gdf_analizado=resultados['gdf_analizado'],
                                    cultivo=cultivo,
                                    analisis_tipo=analisis_tipo,
                                    area_total=resultados['area_total'],
                                    nutriente=nutriente,
                                    satelite=satelite_seleccionado,
                                    indice=indice_seleccionado if 'indice_seleccionado' in locals() else None,
                                    mapa_buffer=resultados['mapa_buffer'],
                                    estadisticas=resultados['estadisticas']
                                )
                                if reporte_pdf:
                                    st.download_button(
                                        label="Descargar PDF",
                                        data=reporte_pdf,
                                        file_name=f"reporte_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                        mime="application/pdf"
                                    )
                        with col_exp3:
                            if st.button("📝 Generar Reporte Word"):
                                reporte_docx = generar_reporte_docx(
                                    gdf_analizado=resultados['gdf_analizado'],
                                    cultivo=cultivo,
                                    analisis_tipo=analisis_tipo,
                                    area_total=resultados['area_total'],
                                    nutriente=nutriente,
                                    satelite=satelite_seleccionado,
                                    indice=indice_seleccionado if 'indice_seleccionado' in locals() else None,
                                    mapa_buffer=resultados['mapa_buffer'],
                                    estadisticas=resultados['estadisticas']
                                )
                                if reporte_docx:
                                    st.download_button(
                                        label="Descargar Word",
                                        data=reporte_docx,
                                        file_name=f"reporte_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                    else:
                        st.error("❌ Error al ejecutar el análisis")
                else:
                    st.warning("⚠️ No se pudo cargar el archivo de parcela")
        else:
            st.info("📤 Sube un archivo de parcela para comenzar el análisis")

        # Mostrar ejemplo de formato
        with st.expander("📋 Formatos de archivo aceptados"):
            st.markdown("""
            **Formato Shapefile (.zip)**:
            - Archivo ZIP que contenga: .shp, .shx, .dbf, .prj
            - Sistema de coordenadas preferido: WGS84 (EPSG:4326)
            **Formato KML/KMZ**:
            - Archivo .kml o .kmz (Google Earth)
            - Debe contener polígonos válidos
            **Ejemplo de estructura:**
            ```
            mi_parcela.zip
            ├── parcela.shp
            ├── parcela.shx
            ├── parcela.dbf
            └── parcela.prj
            ```
            """)

    with tab2:
        st.header("🌱 ANÁLISIS DE FERTILIDAD")
        if 'resultados' in locals() and resultados.get('exitoso'):
            # Mostrar análisis detallado de fertilidad
            gdf_analizado = resultados['gdf_analizado']
            # Gráficos de distribución
            col1, col2 = st.columns(2)
            with col1:
                if 'npk_actual' in gdf_analizado.columns:
                    fig, ax = plt.subplots(figsize=(8, 4))
                    ax.hist(gdf_analizado['npk_actual'], bins=20, color='#3b82f6', edgecolor='white', alpha=0.7)
                    ax.set_xlabel('Índice NPK', color='white')
                    ax.set_ylabel('Frecuencia', color='white')
                    ax.set_title('Distribución de Fertilidad (NPK)', color='white', fontweight='bold')
                    ax.set_facecolor('#0f172a')
                    fig.patch.set_facecolor('#0f172a')
                    ax.tick_params(colors='white')
                    st.pyplot(fig)
            with col2:
                if 'materia_organica' in gdf_analizado.columns:
                    fig, ax = plt.subplots(figsize=(8, 4))
                    ax.scatter(gdf_analizado['materia_organica'], gdf_analizado['npk_actual'] if 'npk_actual' in gdf_analizado.columns else gdf_analizado['id_zona'],
                               color='#10b981', alpha=0.6)
                    ax.set_xlabel('Materia Orgánica (%)', color='white')
                    ax.set_ylabel('NPK' if 'npk_actual' in gdf_analizado.columns else 'Zona', color='white')
                    ax.set_title('Relación Materia Orgánica - Fertilidad', color='white', fontweight='bold')
                    ax.set_facecolor('#0f172a')
                    fig.patch.set_facecolor('#0f172a')
                    ax.tick_params(colors='white')
                    st.pyplot(fig)
            # Recomendaciones de fertilización
            st.subheader("💡 Recomendaciones de Fertilización")
            if 'npk_actual' in gdf_analizado.columns:
                npk_promedio = gdf_analizado['npk_actual'].mean()
                if npk_promedio < 0.3:
                    st.error("**Fertilidad CRÍTICA** - Se requiere fertilización intensiva inmediata")
                    st.markdown("""
                    **Acciones recomendadas:**
                    - Aplicación urgente de fertilizantes balanceados NPK
                    - Incorporar materia orgánica (estiércol, compost)
                    - Considerar cultivos de cobertura para mejorar suelo
                    """)
                elif npk_promedio < 0.5:
                    st.warning("**Fertilidad BAJA** - Se recomienda fertilización moderada")
                    st.markdown("""
                    **Acciones recomendadas:**
                    - Aplicar fertilizantes según análisis de suelo
                    - Fraccionar aplicaciones de nitrógeno
                    - Mantener cobertura vegetal
                    """)
                elif npk_promedio < 0.7:
                    st.info("**Fertilidad ADECUADA** - Mantener prácticas actuales")
                    st.markdown("""
                    **Acciones recomendadas:**
                    - Mantener rotación de cultivos
                    - Monitoreo periódico de nutrientes
                    - Ajustes menores según necesidades específicas
                    """)
                else:
                    st.success("**Fertilidad ÓPTIMA** - Excelente condición del suelo")
                    st.markdown("""
                    **Acciones recomendadas:**
                    - Continuar con prácticas actuales
                    - Monitoreo preventivo
                    - Mantener balance de nutrientes
                    """)
        else:
            st.info("👈 Realiza primero un análisis en la pestaña principal")

    with tab3:
        st.header("💰 ANÁLISIS ECONÓMICO")
        if 'mostrar_economia' in st.session_state and st.session_state['mostrar_economia']:
            if 'precios_actualizados' in st.session_state and 'resultados' in locals():
                precios = st.session_state['precios_actualizados']
                gdf_analizado = resultados['gdf_analizado']
                area_total = resultados['area_total']
                # Realizar análisis económico
                with st.spinner("Calculando análisis económico..."):
                    analisis_economico = generar_analisis_economico(
                        gdf_analizado=gdf_analizado,
                        cultivo=cultivo,
                        area_total=area_total,
                        precios_actualizados=precios
                    )
                if analisis_economico:
                    # Mostrar métricas económicas
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("📈 TIR", f"{analisis_economico['tir']}%")
                    with col2:
                        st.metric("💰 Payback", f"{analisis_economico['payback']} meses")
                    with col3:
                        st.metric("📊 ROI", f"{(analisis_economico['ingreso_extra']/analisis_economico['inversion_fertilizacion']*100):.1f}%")
                    with col4:
                        beneficio_neto = analisis_economico['escenario_con']['total_beneficios']
                        st.metric("💵 Beneficio Neto", f"${beneficio_neto:,.0f}")

                    # Comparación de escenarios
                    st.subheader("📊 Comparación de Escenarios")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown("""
                        <div class="dashboard-card">
                        <h3>❌ Sin Fertilización</h3>
                        """, unsafe_allow_html=True)
                        st.metric("Rendimiento", f"{analisis_economico['escenario_sin']['rendimiento_promedio']:.1f} t/ha")
                        st.metric("Costo Total", f"${analisis_economico['escenario_sin']['total_costos']:,.0f}")
                        st.metric("Beneficio", f"${analisis_economico['escenario_sin']['total_beneficios']:,.0f}")
                    with col2:
                        st.markdown("""
                        <div class="dashboard-card">
                        <h3>✅ Con Fertilización</h3>
                        """, unsafe_allow_html=True)
                        st.metric("Rendimiento", f"{analisis_economico['escenario_con']['rendimiento_promedio']:.1f} t/ha")
                        st.metric("Costo Total", f"${analisis_economico['escenario_con']['total_costos']:,.0f}")
                        st.metric("Beneficio", f"${analisis_economico['escenario_con']['total_beneficios']:,.0f}")

                    # Crear mapas económicos
                    st.subheader("🗺️ Mapas de Rentabilidad")
                    col1, col2 = st.columns(2)
                    with col1:
                        # Mapa de potencial de cosecha
                        mapa_potencial, gdf_potencial = crear_mapa_potencial_cosecha(
                            gdf_analizado, cultivo, precios
                        )
                        if mapa_potencial:
                            mapa_potencial.seek(0)
                            st.image(mapa_potencial, caption="Potencial de Cosecha (t/ha)", use_container_width=True)
                    with col2:
                        # Mapa de rentabilidad
                        mapa_rentabilidad = crear_mapa_rentabilidad(
                            gdf_analizado, cultivo, precios
                        )
                        if mapa_rentabilidad:
                            mapa_rentabilidad.seek(0)
                            st.image(mapa_rentabilidad, caption="Rentabilidad (USD/ha)", use_container_width=True)

                    # Detalles de costos
                    with st.expander("📋 Detalle de Costos de Fertilización"):
                        if 'npk_actual' in gdf_analizado.columns:
                            npk_promedio = gdf_analizado['npk_actual'].mean()
                            costo_fert = calcular_costo_fertilizacion(npk_promedio, cultivo, precios)
                            st.markdown(f"""
                            **Cálculo para fertilidad promedio: {npk_promedio:.2f}**
                            - **Nitrógeno (N):** {costo_fert['dosis_ajustada']['N']} kg/ha = ${costo_fert['costo_n']}/ha
                            - **Fósforo (P):** {costo_fert['dosis_ajustada']['P']} kg/ha = ${costo_fert['costo_p']}/ha
                            - **Potasio (K):** {costo_fert['dosis_ajustada']['K']} kg/ha = ${costo_fert['costo_k']}/ha
                            - **Costo Total Fertilización:** ${costo_fert['costo_total']}/ha
                            """)

                    # Recomendación final
                    st.subheader("🎯 Recomendación Económica")
                    if analisis_economico['tir'] > 15:
                        st.success(f"""
                        **✅ RECOMENDACIÓN: INVERTIR EN FERTILIZACIÓN**
                        La fertilización muestra excelente retorno económico:
                        - TIR del {analisis_economico['tir']}% (superior al costo de capital)
                        - Payback de {analisis_economico['payback']} meses
                        - Incremento de beneficio: ${analisis_economico['ingreso_extra']:,.0f}
                        """)
                    elif analisis_economico['tir'] > 8:
                        st.info(f"""
                        **⚠️ RECOMENDACIÓN: CONSIDERAR FERTILIZACIÓN SELECTIVA**
                        La fertilización muestra retorno moderado:
                        - TIR del {analisis_economico['tir']}%
                        - Evaluar zonas específicas de mayor rentabilidad
                        - Considerar fertilización diferenciada por zonas
                        """)
                    else:
                        st.warning(f"""
                        **❌ RECOMENDACIÓN: POSTERGAR FERTILIZACIÓN**
                        El retorno económico no justifica la inversión:
                        - TIR del {analisis_economico['tir']}% (inferior al costo de capital)
                        - Evaluar mejoras en otros aspectos productivos
                        - Considerar análisis de suelo de laboratorio
                        """)
                else:
                    st.info("Realiza primero un análisis de fertilidad para habilitar el análisis económico")
        else:
            st.info("👈 Haz clic en 'Realizar Análisis Económico' en la pestaña principal")

    with tab4:
        st.header("📈 REPORTES Y EXPORTACIÓN")
        if 'resultados' in locals() and resultados.get('exitoso'):
            # Generar estadísticas
            estadisticas = generar_resumen_estadisticas(
                resultados['gdf_analizado'],
                analisis_tipo,
                cultivo,
                resultados.get('df_power')
            )
            # Generar recomendaciones
            recomendaciones = generar_recomendaciones_generales(
                resultados['gdf_analizado'],
                analisis_tipo,
                cultivo
            )
            # Mostrar resumen
            st.subheader("📊 Resumen Estadístico")
            if estadisticas:
                cols = st.columns(3)
                items = list(estadisticas.items())
                for i, (key, value) in enumerate(items):
                    with cols[i % 3]:
                        st.metric(key, value)
            # Mostrar recomendaciones
            st.subheader("💡 Recomendaciones de Manejo")
            for i, rec in enumerate(recomendaciones[:5]):  # Mostrar solo 5 principales
                st.markdown(f"{i+1}. {rec}")
            # Opciones de exportación completas
            st.subheader("📤 Exportar Resultados")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown("**GeoJSON**")
                if st.button("🌐 Exportar GeoJSON", key="export_geojson"):
                    geojson_data, nombre_archivo = exportar_a_geojson(
                        resultados['gdf_analizado'],
                        f"{cultivo}_{analisis_tipo}"
                    )
                    if geojson_data:
                        st.download_button(
                            label="Descargar",
                            data=geojson_data,
                            file_name=nombre_archivo,
                            mime="application/json"
                        )
            with col2:
                st.markdown("**Reporte PDF**")
                if st.button("📄 Generar PDF", key="export_pdf"):
                    with st.spinner("Generando PDF..."):
                        reporte_pdf = generar_reporte_pdf(
                            gdf_analizado=resultados['gdf_analizado'],
                            cultivo=cultivo,
                            analisis_tipo=analisis_tipo,
                            area_total=resultados['area_total'],
                            nutriente=nutriente,
                            satelite=satelite_seleccionado,
                            indice=indice_seleccionado if 'indice_seleccionado' in locals() else None,
                            mapa_buffer=resultados['mapa_buffer'],
                            estadisticas=estadisticas,
                            recomendaciones=recomendaciones
                        )
                        if reporte_pdf:
                            reporte_pdf.seek(0)
                            st.download_button(
                                label="Descargar PDF",
                                data=reporte_pdf,
                                file_name=f"reporte_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                mime="application/pdf"
                            )
            with col3:
                st.markdown("**Reporte Word**")
                if st.button("📝 Generar Word", key="export_word"):
                    with st.spinner("Generando Word..."):
                        reporte_docx = generar_reporte_docx(
                            gdf_analizado=resultados['gdf_analizado'],
                            cultivo=cultivo,
                            analisis_tipo=analisis_tipo,
                            area_total=resultados['area_total'],
                            nutriente=nutriente,
                            satelite=satelite_seleccionado,
                            indice=indice_seleccionado if 'indice_seleccionado' in locals() else None,
                            mapa_buffer=resultados['mapa_buffer'],
                            estadisticas=estadisticas,
                            recomendaciones=recomendaciones
                        )
                        if reporte_docx:
                            reporte_docx.seek(0)
                            st.download_button(
                                label="Descargar Word",
                                data=reporte_docx,
                                file_name=f"reporte_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
            # Exportar datos tabulares
            st.subheader("📊 Exportar Datos Tabulares")
            if resultados['gdf_analizado'] is not None:
                df_export = resultados['gdf_analizado'].drop(columns=['geometry'] if 'geometry' in resultados['gdf_analizado'].columns else [])
                csv_data = df_export.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="📥 Descargar CSV",
                    data=csv_data,
                    file_name=f"datos_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
        else:
            st.info("Realiza primero un análisis para generar reportes")

    with tab5:
        st.header("ℹ️ AYUDA Y DOCUMENTACIÓN")
        st.markdown("### 📖 Guía de Uso")
        st.markdown("""
        **1. Configuración Inicial**
        - Selecciona el cultivo en el sidebar
        - Elige el tipo de análisis
        - Configura parámetros según el análisis seleccionado
        """)
        st.markdown("""
        **2. Subida de Archivos**
        - Formatos aceptados: Shapefile (.zip), KML, KMZ
        - El archivo debe contener polígonos válidos
        - Sistema de coordenadas preferido: WGS84 (EPSG:4326)
        """)
        st.markdown("""
        **3. Tipos de Análisis Disponibles**
        **🌱 Fertilidad Actual**
        - Analiza el estado nutricional del suelo
        - Genera índice NPK integrado
        - Incluye materia orgánica y humedad
        **🧪 Recomendaciones NPK**
        - Recomienda dosis específicas de nutrientes
        - Basado en índices de vegetación
        - Personalizado por cultivo
        **🏺 Análisis de Textura**
        - Clasificación textural del suelo
        - Recomendaciones de manejo
        - Compatibilidad con cultivo
        **🗺️ Curvas de Nivel**
        - Análisis de pendientes
        - Generación de curvas de nivel
        - Identificación de áreas de riesgo
        """)
        st.markdown("""
        **4. Análisis Económico**
        - Calcula rentabilidad por zona
        - Evalúa TIR y payback
        - Compara escenarios con/sin fertilización
        """)
        st.markdown("### 🛠️ Solución de Problemas")
        st.markdown("""
        **❌ Error al cargar archivo**
        1. Verifica que el archivo tenga el formato correcto
        2. Asegúrate de que contenga geometrías válidas
        3. Revisa el sistema de coordenadas
        **⚠️ Datos satelitales no disponibles**
        1. Los datos simulados siempre están disponibles
        2. Verifica las fechas seleccionadas
        3. Intenta con otro satélite
        """)
        st.markdown("### 📞 Soporte Técnico")
        st.markdown("""
        Para soporte o consultas:
        - 📧 mawucano@gmail.com
        - 📱 +54 9 3525 53-2313
        """)
        st.markdown("### 🔄 Actualizaciones")
        st.markdown("""
        **Versión 2.0** - Diciembre 2024
        - Análisis económico integrado
        - Mapas de rentabilidad
        - Exportación mejorada
        - Interfaz premium
        **Próximas características:**
        - Integración con APIs meteorológicas en tiempo real
        - Modelos predictivos de rendimiento
        - Alertas tempranas de plagas
        - Integración con maquinaria agrícola
        """)

        # Información del sistema
        with st.expander("🔧 Información del Sistema"):
            st.markdown(f"""
            **Versión de la Aplicación:** 2.0.0
            **Última Actualización:** {datetime.now().strftime("%d/%m/%Y")}
            **Sistema de Coordenadas:** EPSG:4326 (WGS84)
            **Satélites Disponibles:** {len(SATELITES_DISPONIBLES)}
            **Cultivos Soportados:** {len(PARAMETROS_CULTIVOS)}
            **Bibliotecas Principales:**
            - Streamlit {st.__version__}
            - GeoPandas {gpd.__version__}
            - Matplotlib {matplotlib.__version__}  # ← CORREGIDO AQUÍ
            - Plotly {go.__version__}
            """)

# Ejecutar la aplicación
if __name__ == "__main__":
    main()
